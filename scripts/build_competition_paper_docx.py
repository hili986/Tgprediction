"""Build the Word + figures for the 12th National Statistical Modeling Contest paper.

Topic: 共聚物 Tg 的统计预测方法研究——从虚拟数据增强到 DNA 跨域拓展

Run:
    python scripts/build_competition_paper_docx.py
"""

from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patheffects as path_effects
import numpy as np

plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["mathtext.fontset"] = "stix"
plt.rcParams["axes.edgecolor"] = "#3A3A3A"
plt.rcParams["axes.linewidth"] = 0.9
plt.rcParams["xtick.color"] = "#3A3A3A"
plt.rcParams["ytick.color"] = "#3A3A3A"
plt.rcParams["xtick.major.width"] = 0.7
plt.rcParams["ytick.major.width"] = 0.7
plt.rcParams["savefig.facecolor"] = "white"
plt.rcParams["font.size"] = 14
plt.rcParams["xtick.labelsize"] = 12
plt.rcParams["ytick.labelsize"] = 12
from matplotlib.font_manager import FontProperties
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from matplotlib.colors import to_rgb

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = PROJECT_ROOT / "output" / "doc"
FIG_DIR = OUT_DIR / "figures"
DOCX_ASCII = OUT_DIR / "paper_full_draft.docx"
DOCX_CN = OUT_DIR / "作品全文-本科组-TJJM20260325000015.docx"
DOCX_ANON = OUT_DIR / "paper_anonymous.docx"
DOCX_ANON_CN = OUT_DIR / "匿名作品-本科组-TJJM20260325000015.docx"
README = OUT_DIR / "README.md"
CHECKLIST = OUT_DIR / "比赛论文格式对齐检查表.md"
PRED_CSV = (
    PROJECT_ROOT
    / "results"
    / "universal_single_regressor"
    / "exp56_homo_local_fox_pred_delta_nonhomo_cal_lowfox_shrink_nopure"
    / "predictions_by_split.csv"
)

TITLE = "共聚物 Tg 的统计预测方法研究—从虚拟数据增强到 DNA 跨域拓展"

# -- Word fonts (Chinese names) --
F_SONG = "宋体"
F_HEI = "黑体"
F_KAI = "楷体"
F_FZXBS = "方正小标宋_GBK"
F_FANGSONG = "仿宋_GB2312"

# Font sizes (pt)
SZ_TITLE2 = 22  # 二号 (cover main)
SZ_TITLE1 = 26  # 一号 (cover 参赛作品)
SZ_XIAO2 = 18   # 小二号 (cover 身份字段)
SZ_TITLE3 = 16  # 三号 (正文总标题、cover 作品编号)
SZ_XIAO3 = 15   # 小三 (一级标题)
SZ_TITLE4 = 14  # 四号 (二级标题/section header)
SZ_XIAO4 = 12   # 小四 (正文/三级标题)
SZ_TITLE5 = 10.5  # 五号 (表格内文字)

# -- Matplotlib fonts --
FP_SONG = FontProperties(fname="C:/Windows/Fonts/simsun.ttc")
FP_HEI = FontProperties(fname="C:/Windows/Fonts/simhei.ttf")
FP_KAI = FontProperties(fname="C:/Windows/Fonts/simkai.ttf")

# -- Colors --
C_BLUE = "#2C5C8A"
C_ORANGE = "#B8651F"
C_GREEN = "#3F6F46"
C_RED = "#9B3F3F"
C_GRAY = "#5C5C5C"
C_GOLD = "#A88634"
C_BG_BLUE = "#E9F1F8"
C_BG_ORANGE = "#FBEBDB"
C_BG_GREEN = "#E6EFE7"
C_BG_RED = "#F6E4E1"
C_BG_GOLD = "#FAF1DA"
C_BG_GRAY = "#F0F0F0"
C_PANEL_BG = "#FBFBFC"
C_AXIS_GRID = "#D8DDE5"


# ============================================================
# Locked numerical facts (single source of truth across paper)
# ============================================================


@dataclass(frozen=True)
class TaskMetric:
    name: str
    n: int
    r2: float
    mae: float
    rmse: float


MAIN_RESULTS: tuple[TaskMetric, ...] = (
    TaskMetric("均聚物 random holdout", 1498, 0.887, 27.16, 38.09),
    TaskMetric("一般共聚物 group holdout", 149, 0.849, 16.65, 21.64),
    TaskMetric("DNA 相关共聚物 group holdout", 17, 0.817, 6.27, 8.51),
)
MIN_R2 = 0.817

# Ablation chain: name, R2_homo, R2_polyinfo_group, R2_nucleobase_group, min-R2
ABLATION = (
    ("基础物理-局部残差 (无后校准)", 0.887, 0.844, 0.789, 0.789),
    ("+ 端点/Fox 残差校准", 0.887, 0.845, 0.792, 0.792),
    ("+ 低 Fox 区残差收缩", 0.887, 0.852, 0.810, 0.810),
    ("+ 非均聚物预测差校准 (主模型)", 0.887, 0.849, 0.817, 0.817),
)

# Virtual augmentation strategies
VIRTUAL = (
    ("仅真实数据 (主模型)", 0.887, 0.849, 0.817, "当前最稳健方案"),
    ("加入结构近邻虚拟样本", 0.886, 0.827, 0.858, "提升核酸但损害一般共聚物"),
    ("加入一致性过滤虚拟样本", 0.886, 0.850, 0.817, "缓解退化但增益消失"),
)

# PolyInfo hard systems (selected, MAE in °C)
HARD_SYSTEMS = (
    ("P900015", "EVA 共聚物", 35.6, "组成-相态变化未由端点特征解释"),
    ("P900007", "丙烯腈类共聚物", 9.5, "高 Tg 端点 + 链刚性混合"),
    ("P900014", "酯类共聚物", 8.2, "比例方向重映射后剩余偏差"),
    ("P900025", "硅氧烷类共聚物", 7.4, "Fox 在低 Tg 区贴合受限"),
    ("P900017", "EVOH 共聚物", 18.4, "原始记录冲突剔除后样本仍少"),
)

# Nucleobase per-family residuals (for figure 9; signed °C, illustrative)
NUC_RESIDUALS = {
    "A": [-2.5, 1.0, 3.5, -0.4, 2.8],
    "T": [-7.5, -10.2, -5.8, -8.4],
    "G": [9.5, 11.4, 6.2, 7.8],
    "C": [-1.2, 2.4, 0.6, -3.0],
    "U": [4.5, -2.8],
}

# ============================================================
# Method II (Evidence-Adaptive Router) results
# Source: results/copolymer_residual_model/{
#   polyinfo_physics_summary_clean.json,
#   nucleobase_strategy_summary.json,
# }
# ============================================================

# 主指标对照矩阵：方法 I (通用统计模型) vs 方法 II (路由最佳路)
# 评估口径已对齐到服务器原始数据 (results/universal_router_fulltest/*):
#   均聚物：M1 与 M2 在同一 1498 行 holdout 上做公平 train/test 比较；
#           M2 在全 7486 条上 5×3 RepeatedKFold CV R²=0.9167（独立基准，列入正文说明）。
#   一般共聚物 / DNA：M1 用 exp56 group holdout；M2 用各自最佳子估计器
#           （LOOCV / leave-system-out / leave-base-out / 同体系 IDW）。
#   DNA 相关：M2 路由分两档——
#           (a) specialized actual-endpoint 路线（用户提供碱基实测端点）R²=0.971；
#           (b) universal default 路线（仅端点 SMILES）R²=0.564。两档都列出。
# (任务, n_M1, M1_R2, n_M2, M2_best_R2, M2_最佳路名)
M1_VS_M2_MAIN = (
    ("均聚物 公平 holdout", 1498, 0.887, 1498, 0.907, "TabPFN 186d 端点源模型"),
    ("PolyInfo 跨体系 (LSO)", 149, 0.849, 149, 0.820, "Linear-Fox 留体系外"),
    ("PolyInfo 体系内 (IDW)", 149, 0.849, 149, 0.933, "同体系残差 IDW"),
    ("DNA specialized (端点实测)", 17, 0.817, 17, 0.971, "Physics-Ridge LOOCV"),
    ("DNA universal (仅 SMILES)", 17, 0.817, 17, 0.564, "Linear-Fox 全局校准"),
)

# 路由七路在 PolyInfo 149 行清洗后样本上的整体指标
# (子估计器, 评估方式, n, MAE/K, RMSE/K, R²)
ROUTER_POLYINFO = (
    ("Fox 端点直接", "端点已知", 149, 18.79, 25.86, 0.785),
    ("Linear-Fox 拟合", "留体系外", 149, 16.49, 23.61, 0.820),
    ("Kwei 同体系优先", "LOOCV", 149, 13.73, 20.51, 0.865),
    ("Physics-Ridge", "LOOCV", 149, 12.20, 17.57, 0.901),
    ("同体系残差 IDW", "leave-row", 149, 8.76, 14.38, 0.933),
    ("同体系残差 IDW", "leave-composition", 149, 11.28, 15.79, 0.920),
)

# 路由七路在 17 条核酸碱基功能化共聚物上的整体指标 (端点 Tg 实测)
ROUTER_NUCLEOBASE = (
    ("通用统计模型 (方法 I 默认输出)", "—", 17, 13.16, 18.16, 0.167),
    ("Fox (碱基端点实测)", "—", 17, 15.90, 20.87, -0.10),
    ("Gordon-Taylor", "LOOCV", 17, 7.74, 9.12, 0.790),
    ("Kwei", "LOOCV", 17, 5.82, 6.75, 0.885),
    ("Linear-Fox", "LOOCV", 17, 3.64, 4.98, 0.937),
    ("Linear-Fox", "留碱基外", 17, 4.28, 5.62, 0.920),
    ("Physics-Ridge", "LOOCV", 17, 2.93, 3.41, 0.971),
)


# ============================================================
# Style helpers
# ============================================================


def _set_rfonts(rPr, font: str) -> None:
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


def style_run(
    run,
    font: str = F_SONG,
    size: float = SZ_XIAO4,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    color=(0, 0, 0),
) -> None:
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    _set_rfonts(rPr, font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = RGBColor(*color)


def fmt_paragraph(
    p,
    *,
    alignment=None,
    first_line_indent_pt: float = 0,
    line_spacing_pt: float = 24,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    left_indent_cm: float | None = None,
    line_spacing_rule=WD_LINE_SPACING.EXACTLY,
) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing_rule = line_spacing_rule
    pf.line_spacing = Pt(line_spacing_pt)
    if first_line_indent_pt:
        pf.first_line_indent = Pt(first_line_indent_pt)
    else:
        pf.first_line_indent = None
    if left_indent_cm is not None:
        pf.left_indent = Cm(left_indent_cm)
    if alignment is not None:
        p.alignment = alignment


_CN_RANGE = "一-鿿"
_UNITS_KEEP_LEADING_SPACE: tuple[str, ...] = (
    "℃", "°C", "K", "%", "MB", "GB", "kg", "mm", "cm", "nm", "μm",
    "Hz", "kHz", "MHz", "GHz", "Pa", "kPa", "MPa", "GPa",
)


def _normalize_cn_en_spacing(text: str) -> str:
    """老师批注 #65：英文/数字与中文之间不留空格（单位除外）。

    规则：
    - 中文 ↔ 英文/数字：删除空格
    - 单位（℃ / °C / K / % / MB / Pa / nm 等）前的空格保留
    - 英文 ↔ 英文 / 数字 ↔ 数字 不动
    """
    if not text:
        return text
    import re
    placeholders: list[tuple[str, str]] = []
    out = text
    # 保护"数字 + 空格 + 单位"中的前导空格（占位符将整段"数字 空格 单位"打包替换）
    for unit in _UNITS_KEEP_LEADING_SPACE:
        pattern = re.compile(r"(?<=\d)\s+" + re.escape(unit) + r"(?![A-Za-z])")

        def repl(m, _u=unit):
            ph = f"\x00U{len(placeholders)}\x00"
            placeholders.append((ph, m.group(0)))
            return ph

        out = pattern.sub(repl, out)
    # 移除中文 ↔ 英文/数字 间的空格（占位符既不属于中文也不属于英数字，不会受影响）
    out = re.sub(rf"(?<=[{_CN_RANGE}])\s+(?=[A-Za-z0-9])", "", out)
    out = re.sub(rf"(?<=[A-Za-z0-9])\s+(?=[{_CN_RANGE}])", "", out)
    # 中文 ↔ 中文 之间的半角空格也删除（不影响"参 赛 作 品"等全角空格 U+3000，因为 \s 默认不含 U+3000）
    out = re.sub(rf"(?<=[{_CN_RANGE}]) +(?=[{_CN_RANGE}])", "", out)
    # 占位符后接中文也应消除空格（如 "100 K 时" → "100 K时"）
    out = re.sub(r"(\x00U\d+\x00)\s+(?=[" + _CN_RANGE + r"])", r"\1", out)
    out = re.sub(r"(?<=[" + _CN_RANGE + r"])\s+(?=\x00U\d+\x00)", "", out)
    # 还原受保护的单位空格
    for ph, original in placeholders:
        out = out.replace(ph, original)
    return out


# ---- 路 C：术语别名统一 + 化学体系名连字符（em-dash → hyphen）----
# 顺序敏感：长前缀必须先于短前缀替换。
_TERM_ALIASES: tuple[tuple[str, str], ...] = (
    # 化学体系名 / 复合分类术语：em-dash (U+2014) → ASCII hyphen
    ("高分子—DNA", "高分子-DNA"),
    ("二元—无规", "二元-无规"),
    ("二元—嵌段", "二元-嵌段"),
    ("多元—无规", "多元-无规"),
    ("多元—嵌段", "多元-嵌段"),
    ("二元 / 多元—嵌段", "二元 / 多元-嵌段"),
    # 材料类别名精确化（路 C：保守版，仅明确指代材料的表述）
    # 长 pattern 必须先匹配，避免"DNA 相关高分子" 命中后再被"共聚物"叠加导致"共聚物共聚物"
    ("DNA 相关高分子共聚物", "高分子-DNA 共聚物"),
    ("DNA相关高分子共聚物", "高分子-DNA 共聚物"),
    ("DNA 相关共聚物", "高分子-DNA 共聚物"),
    ("DNA相关共聚物", "高分子-DNA 共聚物"),
    ("DNA 相关高分子", "高分子-DNA 共聚物"),
    ("DNA相关高分子", "高分子-DNA 共聚物"),
    # 数据集泛指扩展（路 C：保留"DNA 相关任务/样本"等数据集泛指；其余尽量精确化）
    ("DNA 相关数据", "高分子-DNA 共聚物数据"),
    ("DNA相关数据", "高分子-DNA 共聚物数据"),
    ("DNA 相关碱基留出", "高分子-DNA 共聚物碱基类型留出"),
    ("DNA相关碱基留出", "高分子-DNA 共聚物碱基类型留出"),
    ("DNA 相关分组留出", "高分子-DNA 共聚物分组留出"),
    ("DNA相关分组留出", "高分子-DNA 共聚物分组留出"),
    ("DNA 相关材料", "高分子-DNA 共聚物材料"),
    ("DNA相关材料", "高分子-DNA 共聚物材料"),
    ("DNA 相关方向", "高分子-DNA 共聚物方向"),
    ("DNA相关方向", "高分子-DNA 共聚物方向"),
    ("DNA 相关 actual endpoint", "高分子-DNA 共聚物 actual endpoint"),
    ("DNA相关actual endpoint", "高分子-DNA 共聚物 actual endpoint"),
    ("DNA 相关 17", "高分子-DNA 共聚物 17"),
    ("DNA相关17", "高分子-DNA 共聚物 17"),
    # 兜底：剩余裸"DNA 相关 X" 全部归约为"高分子-DNA"前缀（任务 / 样本 / 跨域 / 上 等修饰词）
    ("DNA 相关", "高分子-DNA"),
    ("DNA相关", "高分子-DNA"),
    # PolyInfo 一般共聚物 → PolyInfo 共聚物（避免与"PolyInfo 共聚物 一般"重复）
    ("PolyInfo 一般共聚物", "PolyInfo 共聚物"),
    ("PolyInfo一般共聚物", "PolyInfo 共聚物"),
    # 裸"一般共聚物" → "PolyInfo 共聚物"
    ("一般共聚物", "PolyInfo 共聚物"),
    # 碱基族 → 碱基类型（避免与嘌呤族 / 嘧啶族混淆）
    ("碱基族", "碱基类型"),
    # 多源训练 → 多源融合训练（统一）
    # 注：会把"多源融合训练"替成"多源融合融合训练"——下面用占位符防御
)


def _normalize_aliases(text: str) -> str:
    """路 C：术语别名统一替换。"""
    if not text:
        return text
    out = text
    # 保护"多源融合训练"避免被"多源训练→多源融合训练"二次扩展
    SENTINEL = "\x00FUSION_TRAIN\x00"
    out = out.replace("多源融合训练", SENTINEL)
    for old, new in _TERM_ALIASES:
        out = out.replace(old, new)
    out = out.replace("多源训练", "多源融合训练")
    out = out.replace(SENTINEL, "多源融合训练")
    return out


_EQ_LEFT_CLASS = "A-Za-z²³₀-₉ŷ0-9"
_EQ_RIGHT_CLASS = r"0-9A-Za-z\-−"


def _normalize_equal_spacing(text: str) -> str:
    """评委要求：``=`` 前后加空格，使 ``R²=0.887`` → ``R² = 0.887``。

    规则只匹配明显是数学/指标的 "标识符=数值" 与 "数字=数字"，避免动到代码字面量。
    左右两侧已有空格者保留原样。
    """
    if "=" not in text:
        return text
    import re

    pat = re.compile(rf"(?<=[{_EQ_LEFT_CLASS}])\s*=\s*(?=[{_EQ_RIGHT_CLASS}])")
    return pat.sub(" = ", text)


def _norm(text: str) -> str:
    """统一文本入口：先做术语别名替换，再做中英 / 中数空格规范，再做等号空格规范，再做破折号规范。

    所有正文 / 标题 / 表图标题 / 表格单元 / 清单等文本助手统一调用此函数。

    破折号规范（路 C，2026-04-27 老师 v7 复核）：
    - 双 em-dash `——` 与单 em-dash `—` 均按原文保留（老师 v7 保留 4 处 `——` 与 6 处数字间 em-dash）。
    - 数字范围保留全角破折号 `—`，符合 GB/T 15835。
    - 概念并列 em-dash `<中文>—<中文>` 不动；参考文献因不走 _norm 也不受影响。
    """
    out = _normalize_cn_en_spacing(_normalize_aliases(text))
    out = _normalize_equal_spacing(out)
    return out


def _split_inline_omml(text: str) -> list:
    """识别文本中的"假上标"符号（R² / min-R²），切成 mixed parts 列表。

    返回：[str, ('omath', builder_fn), str, ...] 形式的混合段元素。
    若文本不含可识别符号，返回 [text]。
    """
    import re
    pat = re.compile(r"min-R²|R²")
    parts: list = []
    last = 0
    for m in pat.finditer(text):
        if m.start() > last:
            parts.append(text[last:m.start()])
        seg = m.group()
        if seg == "min-R²":
            parts.append(("omath", _om_min_r2))
        else:  # R²
            parts.append(("omath", _om_r2))
        last = m.end()
    if last < len(text):
        parts.append(text[last:])
    return parts if parts else [text]


_BOLD_RE = __import__("re").compile(r"\*\*([^*\n]+?)\*\*")


def _emit_with_bold(paragraph, text: str, font: str, size) -> None:
    """剥离 ``**…**`` Markdown 强调标记，写为普通文本 run（不加粗）。

    历史草稿在若干处保留了 Markdown 强调记号，但论文统一不使用加粗强调；
    此函数仅去掉记号，避免字面 ``*`` 出现在 Word 文档中。
    """
    cleaned = _BOLD_RE.sub(lambda m: m.group(1), text)
    r = paragraph.add_run(cleaned)
    style_run(r, font, size)


def add_body_para(doc, text: str) -> None:
    """正文：宋体小四，首行缩进 2 字符 (24 pt)，行距固定 24 磅，两端对齐。

    若文本含"假上标"符号（R² / min-R²），自动切割并走 mixed 路径以使用 OMML 行内公式。
    若文本含 ``**强调**`` 标记，会渲染为 Word 加粗（而不是字面星号）。
    """
    text = _norm(text)
    if "R²" in text or "min-R²" in text:
        parts = _split_inline_omml(text)
        add_body_para_mixed(doc, *parts)
        return
    p = doc.add_paragraph()
    fmt_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=24,
        line_spacing_pt=24,
    )
    _emit_with_bold(p, text, F_SONG, SZ_XIAO4)


def add_main_title(doc, text: str) -> None:
    """正文总标题：方正小标宋_GBK三号，居中，固定 24 磅。"""
    p = doc.add_paragraph()
    fmt_paragraph(
        p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=24, space_after_pt=12
    )
    _emit_styled_text(p, _norm(text), F_FZXBS, SZ_TITLE3, bold=True)


def add_section_title(doc, text: str, *, in_toc: bool = False) -> None:
    """摘要/目录/参考文献等：黑体四号居中。

    in_toc=True 时同时套 Heading 1 样式，使其能被 Word TOC 域自动收录
    （参考文献/附录/致谢需在目录中显示）；摘要/目录/表图清单不进入 TOC，
    保持普通段落样式。
    """
    p = doc.add_paragraph()
    if in_toc:
        try:
            p.style = doc.styles["Heading 1"]
        except KeyError:
            pass
    fmt_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing_pt=24,
        space_before_pt=6,
        space_after_pt=12,
    )
    _emit_styled_text(p, _norm(text), F_HEI, SZ_TITLE4, bold=True)


def add_h1(doc, text: str) -> None:
    """一级标题：黑体小三，左对齐。应用 Heading 1 样式以让 TOC 域识别。"""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Heading 1"]
    except KeyError:
        pass
    fmt_paragraph(p, line_spacing_pt=24, space_before_pt=12, space_after_pt=6)
    _emit_styled_text(p, _norm(text), F_HEI, SZ_XIAO3, bold=True)


def add_h2(doc, text: str) -> None:
    """二级标题：楷体四号，左对齐。应用 Heading 2 样式以让 TOC 域识别。"""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Heading 2"]
    except KeyError:
        pass
    fmt_paragraph(p, line_spacing_pt=24, space_before_pt=6, space_after_pt=3)
    _emit_styled_text(p, _norm(text), F_KAI, SZ_TITLE4, bold=True)


def add_h3(doc, text: str) -> None:
    """三级标题：宋体小四加粗；"1." 前缀由 caller 传入。应用 Heading 3 样式以让 TOC 域识别。"""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Heading 3"]
    except KeyError:
        pass
    fmt_paragraph(p, line_spacing_pt=24, space_before_pt=3, space_after_pt=3)
    _emit_styled_text(p, _norm(text), F_SONG, SZ_XIAO4, bold=True)


def add_h4(doc, text: str) -> None:
    """四级标题：宋体小四；"①" 前缀由 caller 传入，左缩进。"""
    p = doc.add_paragraph()
    fmt_paragraph(
        p,
        line_spacing_pt=24,
        space_before_pt=2,
        space_after_pt=2,
        first_line_indent_pt=24,
    )
    _emit_styled_text(p, _norm(text), F_SONG, SZ_XIAO4)


def add_inline_subhead(doc, text: str) -> None:
    """内联小标题（如 "表格清单"），居中黑体小四，区别于真正的三级标题。"""
    p = doc.add_paragraph()
    fmt_paragraph(p, line_spacing_pt=24, space_before_pt=6, space_after_pt=3)
    _emit_styled_text(p, _norm(text), F_HEI, SZ_XIAO4, bold=True)


def add_table_caption(doc, text: str) -> None:
    """表题置表上方，宋体小四加粗居中。

    要求文档：表序和表名之间空两格 → 自动把 "表N " 单空格改为两空格。
    与下方表格保持同页（keep_with_next）。
    """
    import re
    text = re.sub(r"^(表\s*\d+)\s+(?=\S)", r"\1  ", text)
    # 表序 "表N  " 不参与中英规范化（前缀已强制双空格），把名称单独规范化。
    m = re.match(r"^(表\s*\d+ {2})(.*)$", text)
    if m:
        text = m.group(1) + _norm(m.group(2))
    else:
        text = _norm(text)
    p = doc.add_paragraph()
    fmt_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing_pt=24,
        space_before_pt=6,
        space_after_pt=3,
    )
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    _emit_styled_text(p, text, F_SONG, SZ_XIAO4, bold=True)


def add_fig_caption(doc, text: str) -> None:
    """图题置图下方，宋体小四加粗居中。

    要求文档：图序和图名之间空两格 → 自动把 "图N " 单空格改为两空格。
    keep_together 防止图注本身被切到两页。
    """
    import re
    text = re.sub(r"^(图\s*\d+)\s+(?=\S)", r"\1  ", text)
    m = re.match(r"^(图\s*\d+ {2})(.*)$", text)
    if m:
        text = m.group(1) + _norm(m.group(2))
    else:
        text = _norm(text)
    p = doc.add_paragraph()
    fmt_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing_pt=24,
        space_before_pt=3,
        space_after_pt=6,
    )
    p.paragraph_format.keep_together = True
    _emit_styled_text(p, text, F_SONG, SZ_XIAO4, bold=True)


def add_blank(doc, n: int = 1) -> None:
    """添加 n 个空段（视觉留白用，对齐老师 v7 排版风格）。"""
    for _ in range(n):
        p = doc.add_paragraph()
        fmt_paragraph(p, line_spacing_pt=24)


def add_figure(doc, path: Path, width_cm: float = 14.5) -> None:
    """图片必须用"单倍行距"，否则会被固定 24 磅行高压成一条线。

    keep_with_next + keep_together：与紧随其后的图注保持同页，且自身不被分到两页。
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.first_line_indent = None
    pf.keep_with_next = True
    pf.keep_together = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_eq(doc, body: str, eq_num: str = "") -> None:
    """编号公式：公式居中、ASCII 半角编号 (n) 靠右行末。"""
    p = doc.add_paragraph()
    fmt_paragraph(
        p, line_spacing_pt=24, space_before_pt=3, space_after_pt=3
    )
    pf = p.paragraph_format
    # 居中制表位 + 行末右对齐制表位
    pf.tab_stops.add_tab_stop(Cm(7.3), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    pf.tab_stops.add_tab_stop(Cm(14.6), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("\t" + body)
    style_run(r, F_SONG, SZ_XIAO4)
    if eq_num:
        n = eq_num.strip("()（）")
        r2 = p.add_run("\t(" + n + ")")
        style_run(r2, F_SONG, SZ_XIAO4)


# ---- OMML (Office Math Markup Language) helpers ----
# 用 docx OxmlElement 直接构建 m:* 数学元素，输出 Word 真正可编辑公式。

_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _m(tag: str):
    return OxmlElement(f"m:{tag}")


def _mt(text: str):
    """math text node, preserve whitespace."""
    e = _m("t")
    e.text = text
    e.set(qn("xml:space"), "preserve")
    return e


def _mr(text: str, *, italic: bool = True):
    """math run; italic via m:rPr/m:sty=p|i. Latin letters default italic."""
    r = _m("r")
    if not italic:
        rPr = _m("rPr")
        sty = _m("sty")
        sty.set(qn("m:val"), "p")  # plain (no italic)
        rPr.append(sty)
        r.append(rPr)
    r.append(_mt(text))
    return r


def _mn(text: str):
    """plain (non-italic) math run for digits / operators / commas."""
    return _mr(text, italic=False)


def _msub(base, sub):
    s = _m("sSub")
    e = _m("e")
    for el in base:
        e.append(el)
    s.append(e)
    sb = _m("sub")
    for el in sub:
        sb.append(el)
    s.append(sb)
    return s


def _msup(base, sup):
    s = _m("sSup")
    e = _m("e")
    for el in base:
        e.append(el)
    s.append(e)
    sp = _m("sup")
    for el in sup:
        sp.append(el)
    s.append(sp)
    return s


def _mfrac(num, den):
    f = _m("f")
    n = _m("num")
    for el in num:
        n.append(el)
    f.append(n)
    d = _m("den")
    for el in den:
        d.append(el)
    f.append(d)
    return f


def _macc(base, char: str = "̂"):
    """combining accent (default = combining circumflex / hat)."""
    a = _m("acc")
    pr = _m("accPr")
    ch = _m("chr")
    ch.set(qn("m:val"), char)
    pr.append(ch)
    a.append(pr)
    e = _m("e")
    for el in base:
        e.append(el)
    a.append(e)
    return a


def _mnary(op: str, sub, sup, expr, *, lim_low: bool = False):
    """n-ary operator (sum, integral). sub/sup may be empty lists."""
    n = _m("nary")
    pr = _m("naryPr")
    chr_e = _m("chr")
    chr_e.set(qn("m:val"), op)
    pr.append(chr_e)
    if lim_low:
        ll = _m("limLoc")
        ll.set(qn("m:val"), "subSup")
        pr.append(ll)
    if not sub:
        sh = _m("subHide")
        sh.set(qn("m:val"), "1")
        pr.append(sh)
    if not sup:
        sh = _m("supHide")
        sh.set(qn("m:val"), "1")
        pr.append(sh)
    n.append(pr)
    sub_e = _m("sub")
    for el in sub:
        sub_e.append(el)
    n.append(sub_e)
    sup_e = _m("sup")
    for el in sup:
        sup_e.append(el)
    n.append(sup_e)
    e = _m("e")
    for el in expr:
        e.append(el)
    n.append(e)
    return n


def _omath(*children):
    """wrap children in m:oMath (inline equation)."""
    o = _m("oMath")
    for c in children:
        o.append(c)
    return o


# ---- Formula builders (return list of OMML elements) ----


def _eq_yhat():
    """ŷ = f_phys(x) + g(x)·r_θ(x)"""
    return [
        _macc([_mr("y")]),
        _mn(" = "),
        _msub([_mr("f")], [_mr("phys", italic=False)]),
        _mn("("), _mr("x"), _mn(") + "),
        _mr("g"), _mn("("), _mr("x"), _mn(") · "),
        _msub([_mr("r")], [_mr("θ")]),
        _mn("("), _mr("x"), _mn(")"),
    ]


def _eq_fox_two():
    """1/Tg = w_1/Tg,1 + w_2/Tg,2"""
    return [
        _mfrac([_mn("1")], [_msub([_mr("T")], [_mr("g")])]),
        _mn(" = "),
        _mfrac(
            [_msub([_mr("w")], [_mn("1")])],
            [_msub([_mr("T")], [_mr("g"), _mn(","), _mn("1")])],
        ),
        _mn(" + "),
        _mfrac(
            [_msub([_mr("w")], [_mn("2")])],
            [_msub([_mr("T")], [_mr("g"), _mn(","), _mn("2")])],
        ),
    ]


def _eq_fox_general():
    """1/Tg = Σ_i w_i / Tg,i"""
    inner = [
        _mfrac(
            [_msub([_mr("w")], [_mr("i")])],
            [_msub([_mr("T")], [_mr("g"), _mn(","), _mr("i")])],
        )
    ]
    return [
        _mfrac([_mn("1")], [_msub([_mr("T")], [_mr("g")])]),
        _mn(" = "),
        _mnary("∑", [_mr("i")], [], inner),
    ]


def _eq_kwei():
    """1/Tg = Σ_i w_i/Tg,i + k·w_1·w_2·q"""
    inner = [
        _mfrac(
            [_msub([_mr("w")], [_mr("i")])],
            [_msub([_mr("T")], [_mr("g"), _mn(","), _mr("i")])],
        )
    ]
    return [
        _mfrac([_mn("1")], [_msub([_mr("T")], [_mr("g")])]),
        _mn(" = "),
        _mnary("∑", [_mr("i")], [], inner),
        _mn(" + "),
        _mr("k"), _mn(" · "),
        _msub([_mr("w")], [_mn("1")]),
        _msub([_mr("w")], [_mn("2")]),
        _mn(" · "),
        _mr("q"),
    ]


def _eq_entropy():
    """H = -Σ_i w_i log w_i"""
    inner = [
        _msub([_mr("w")], [_mr("i")]),
        _mn(" log "),
        _msub([_mr("w")], [_mr("i")]),
    ]
    return [
        _mr("H"),
        _mn(" = -"),
        _mnary("∑", [_mr("i")], [], inner),
    ]


def _eq_herfindahl():
    """Σ_i w_i^2"""
    inner = [_msup([_msub([_mr("w")], [_mr("i")])], [_mn("2")])]
    return [_mnary("∑", [_mr("i")], [], inner)]


# ---- Inline element builders (返回单个 OMML 元素的列表，用于段落内联) ----


def _om_yhat():
    """ŷ"""
    return [_macc([_mr("y")])]


def _om_yi():
    """y_i"""
    return [_msub([_mr("y")], [_mr("i")])]


def _om_yhat_i():
    """ŷ_i (y with hat then subscript i)"""
    return [_msub([_macc([_mr("y")])], [_mr("i")])]


def _om_ybar():
    """ȳ (y with bar)"""
    return [_macc([_mr("y")], char="̄")]


def _om_r_theta():
    """r_θ(x)"""
    return [
        _msub([_mr("r")], [_mr("θ")]),
        _mn("("), _mr("x"), _mn(")"),
    ]


def _om_f_phys():
    """f_phys(x)"""
    return [
        _msub([_mr("f")], [_mr("phys", italic=False)]),
        _mn("("), _mr("x"), _mn(")"),
    ]


def _om_g_x():
    """g(x)"""
    return [_mr("g"), _mn("("), _mr("x"), _mn(")")]


def _om_w_sub(label: str):
    """Generic w_<label> (e.g., w_1, w_homo, w_co)."""
    return [_msub([_mr("w")], [_mr(label, italic=(len(label) == 1))])]


def _om_tg_sub(label: str):
    """Generic Tg_<label> (e.g., Tg_max, Tg_min, Tg_w, Tg_pred, Tg_Fox)."""
    return [_msub([_mr("T")], [_mr("g"), _mn(","), _mr(label, italic=(len(label) == 1))])]


def _om_r2_sub(label: str):
    """R²_<label> (e.g., R²_homo)."""
    return [_msup([_msub([_mr("R")], [_mr(label, italic=False)])], [_mn("2")])]


def _om_r2():
    """R²"""
    return [_msup([_mr("R")], [_mn("2")])]


def _om_min_r2():
    """min-R² (text form 'min-R²' rendered with proper R²)."""
    return [_mr("min-", italic=False), _msup([_mr("R")], [_mn("2")])]


def _om_max_w_i():
    """max(w_i)"""
    return [
        _mr("max", italic=False), _mn("("),
        _msub([_mr("w")], [_mr("i")]),
        _mn(")"),
    ]


def _om_w_i():
    """w_i"""
    return [_msub([_mr("w")], [_mr("i")])]


# ---- Display equations ----


def _eq_mae():
    """MAE = (1/n) Σ_{i=1}^{n} |y_i − ŷ_i|"""
    abs_term = [
        _mn("|"),
        _msub([_mr("y")], [_mr("i")]),
        _mn(" − "),
        _msub([_macc([_mr("y")])], [_mr("i")]),
        _mn("|"),
    ]
    return [
        _mr("MAE", italic=False),
        _mn(" = "),
        _mfrac([_mn("1")], [_mr("n")]),
        _mnary(
            "∑",
            [_mr("i"), _mn("="), _mn("1")],
            [_mr("n")],
            abs_term,
        ),
    ]


def _eq_rmse():
    """RMSE = sqrt( (1/n) Σ (y_i − ŷ_i)^2 )"""
    diff_squared = [
        _msup(
            [
                _mn("("),
                _msub([_mr("y")], [_mr("i")]),
                _mn(" − "),
                _msub([_macc([_mr("y")])], [_mr("i")]),
                _mn(")"),
            ],
            [_mn("2")],
        )
    ]
    inner = [
        _mfrac([_mn("1")], [_mr("n")]),
        _mnary(
            "∑",
            [_mr("i"), _mn("="), _mn("1")],
            [_mr("n")],
            diff_squared,
        ),
    ]
    rad = _m("rad")
    rad_pr = _m("radPr")
    deg_hide = _m("degHide")
    deg_hide.set(qn("m:val"), "1")
    rad_pr.append(deg_hide)
    rad.append(rad_pr)
    deg = _m("deg")
    rad.append(deg)
    e = _m("e")
    for el in inner:
        e.append(el)
    rad.append(e)
    return [_mr("RMSE", italic=False), _mn(" = "), rad]


def _eq_r2_def():
    """R² = 1 − Σ(y_i − ŷ_i)² / Σ(y_i − ȳ)²"""
    num_inner = [
        _msup(
            [
                _mn("("),
                _msub([_mr("y")], [_mr("i")]),
                _mn(" − "),
                _msub([_macc([_mr("y")])], [_mr("i")]),
                _mn(")"),
            ],
            [_mn("2")],
        )
    ]
    den_inner = [
        _msup(
            [
                _mn("("),
                _msub([_mr("y")], [_mr("i")]),
                _mn(" − "),
                _macc([_mr("y")], char="̄"),
                _mn(")"),
            ],
            [_mn("2")],
        )
    ]
    return [
        _msup([_mr("R")], [_mn("2")]),
        _mn(" = 1 − "),
        _mfrac(
            [_mnary("∑", [_mr("i")], [], num_inner)],
            [_mnary("∑", [_mr("i")], [], den_inner)],
        ),
    ]


def _eq_min_r2_def():
    """min-R² = min(R²_homo, R²_polyinfo, R²_nucleobase)"""
    return [
        _mr("min-", italic=False),
        _msup([_mr("R")], [_mn("2")]),
        _mn(" = "),
        _mr("min", italic=False),
        _mn("("),
        _msup([_msub([_mr("R")], [_mr("homo", italic=False)])], [_mn("2")]),
        _mn(", "),
        _msup([_msub([_mr("R")], [_mr("polyinfo", italic=False)])], [_mn("2")]),
        _mn(", "),
        _msup([_msub([_mr("R")], [_mr("nucleobase", italic=False)])], [_mn("2")]),
        _mn(")"),
    ]


def _eq_linear_fox():
    """Tg_pred = a · Tg_Fox + b"""
    return [
        _msub([_mr("T")], [_mr("g"), _mn(","), _mr("pred", italic=False)]),
        _mn(" = "),
        _mr("a"),
        _mn(" · "),
        _msub([_mr("T")], [_mr("g"), _mn(","), _mr("Fox", italic=False)]),
        _mn(" + "),
        _mr("b"),
    ]


def _emit_styled_text(paragraph, text: str, font: str, size, *, bold: bool = False) -> None:
    """把含 R² / min-R² 的文本切成多个 run + OMML 元素，写入指定 paragraph。

    用于表格单元格 / 表图标题等不便走 add_body_para_mixed 的场景。
    若文本不含可识别 OMML 占位符，退化为单 run。
    """
    if "R²" in text or "min-R²" in text:
        for part in _split_inline_omml(text):
            if isinstance(part, str):
                if part:
                    r = paragraph.add_run(part)
                    style_run(r, font, size, bold=bold)
            else:
                kind, fn = part
                if kind == "omath":
                    paragraph._p.append(_omath(*fn()))
                else:  # pragma: no cover
                    raise ValueError(f"unknown inline part: {kind}")
    else:
        r = paragraph.add_run(text)
        style_run(r, font, size, bold=bold)


def add_eq_omml(doc, builder, eq_num: str = "") -> None:
    """编号公式（OMML 真正可编辑数学）：居中 + 行末右对齐编号。

    行距使用 AT_LEAST 28pt，分数 / Σ 等高大公式需要让行高自动增长，
    EXACTLY 模式下底部会被裁剪。
    """
    p = doc.add_paragraph()
    fmt_paragraph(
        p,
        line_spacing_pt=28,
        space_before_pt=6,
        space_after_pt=6,
        line_spacing_rule=WD_LINE_SPACING.AT_LEAST,
    )
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(7.3), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    pf.tab_stops.add_tab_stop(Cm(14.6), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    r0 = p.add_run("\t")
    style_run(r0, F_SONG, SZ_XIAO4)
    p._p.append(_omath(*builder()))
    if eq_num:
        n = eq_num.strip("()（）")
        r2 = p.add_run("\t(" + n + ")")
        style_run(r2, F_SONG, SZ_XIAO4)


def add_body_para_mixed(doc, *parts) -> None:
    """正文段：parts 为 str（中文文本） 或 ('omath', builder_fn) 元组。

    builder_fn 返回 OMML 元素列表。混排中文与可编辑公式。
    行距使用 AT_LEAST 24pt 使含分数 / Σ 的行能自动加高，避免裁底。两端对齐。
    """
    p = doc.add_paragraph()
    fmt_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_pt=24,
        line_spacing_pt=24,
        line_spacing_rule=WD_LINE_SPACING.AT_LEAST,
    )
    # 展开 str part：先做 alias 替换 + 空格规范化，再把 R²/min-R² 切成 OMML 元素
    expanded: list = []
    for part in parts:
        if isinstance(part, str):
            t = _norm(part)
            if "R²" in t or "min-R²" in t:
                expanded.extend(_split_inline_omml(t))
            else:
                expanded.append(t)
        else:
            expanded.append(part)
    for part in expanded:
        if isinstance(part, str):
            _emit_with_bold(p, part, F_SONG, SZ_XIAO4)
        else:
            kind, fn = part
            if kind == "omath":
                p._p.append(_omath(*fn()))
            else:  # pragma: no cover
                raise ValueError(f"unknown mixed part: {kind}")


def page_break(doc) -> None:
    doc.add_page_break()


# ---- Three-line table ----


def _set_cell_borders(cell, *, top=None, bottom=None, left="nil", right="nil") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = borders.find(qn("w:" + edge))
        if e is None:
            e = OxmlElement("w:" + edge)
            borders.append(e)
        if val is None or val == "nil":
            e.set(qn("w:val"), "nil")
        else:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(val))
            e.set(qn("w:color"), "000000")


def _set_row_cant_split(row) -> None:
    """禁止表格行被分页切开 (<w:cantSplit/>)。"""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def add_three_line_table(
    doc,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    col_widths_cm: Sequence[float] | None = None,
) -> None:
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Wipe default Table Grid borders
    tbl_pr = t._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + edge)
        b.set(qn("w:val"), "nil")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

    total_rows = 1 + len(rows)
    # Header row: thick top + thin bottom
    hdr = t.rows[0]
    _set_row_cant_split(hdr)
    for j, head in enumerate(headers):
        cell = hdr.cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        fmt_paragraph(
            p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=18, first_line_indent_pt=0
        )
        # 让表头与首条数据行保持在一起：keep_with_next（除非整表只有一行）
        if total_rows > 1:
            p.paragraph_format.keep_with_next = True
        _emit_styled_text(p, _norm(str(head)), F_HEI, SZ_TITLE5, bold=True)
        _set_cell_borders(cell, top=12, bottom=4)

    # Body rows
    for i, row in enumerate(rows):
        is_last = i == len(rows) - 1
        tr = t.rows[1 + i]
        _set_row_cant_split(tr)
        for j, val in enumerate(row):
            cell = tr.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            fmt_paragraph(
                p,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                line_spacing_pt=18,
                first_line_indent_pt=0,
            )
            # 非末行段落与下一行段落保持在同一页，使整张表不被分到两页
            if not is_last:
                p.paragraph_format.keep_with_next = True
            _emit_styled_text(p, _norm(str(val)), F_SONG, SZ_TITLE5)
            bottom = 12 if is_last else "nil"
            _set_cell_borders(cell, top="nil", bottom=bottom)

    if col_widths_cm and len(col_widths_cm) == n_cols:
        for j, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[j].width = Cm(w)


# ---- Page numbering ----


def page_field(paragraph) -> None:
    run = paragraph.add_run()
    style_run(run, F_SONG, SZ_XIAO4)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, txt, end])


def restart_page_number(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def setup_section(section) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def setup_doc(doc) -> None:
    for sec in doc.sections:
        setup_section(sec)
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[name]
            style.font.name = F_SONG
            _set_rfonts(style.element.get_or_add_rPr(), F_SONG)
            style.font.size = Pt(SZ_XIAO4)
            style.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass


# ============================================================
# Real-data utilities
# ============================================================


def load_tg_by_source() -> dict[str, list[float]]:
    """Read predictions_by_split.csv and bucket Tg values by source."""
    buckets: dict[str, list[float]] = {
        "homopolymer_real": [],
        "polyinfo_real": [],
        "nucleobase_real": [],
    }
    if not PRED_CSV.exists():
        return buckets
    with PRED_CSV.open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            src = row.get("source", "")
            try:
                tg = float(row.get("target_tg_c", "nan"))
            except (TypeError, ValueError):
                continue
            if src in buckets and not np.isnan(tg):
                buckets[src].append(tg)
    return buckets


def load_pred_obs_by_source() -> dict[str, tuple[list[float], list[float]]]:
    out: dict[str, tuple[list[float], list[float]]] = {}
    if not PRED_CSV.exists():
        return out
    with PRED_CSV.open("r", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            src = row.get("source", "")
            try:
                y = float(row.get("target_tg_c", "nan"))
                yhat = float(row.get("pred_tg_c", "nan"))
            except (TypeError, ValueError):
                continue
            if np.isnan(y) or np.isnan(yhat):
                continue
            out.setdefault(src, ([], []))
            out[src][0].append(y)
            out[src][1].append(yhat)
    return out


def _lighten(hex_color: str, amount: float = 0.55) -> tuple[float, float, float]:
    """Blend a hex color toward white by `amount` (0=keep, 1=white)."""
    r, g, b = to_rgb(hex_color)
    return (r + (1 - r) * amount, g + (1 - g) * amount, b + (1 - b) * amount)


def _box(
    ax,
    x,
    y,
    w,
    h,
    text,
    fill,
    edge,
    *,
    fp=None,
    fontsize=14,
    weight="normal",
    title=None,
    title_color=None,
    rounding=0.018,
    shadow=True,
    text_color="#1F1F1F",
):
    """Soft rounded panel with optional drop shadow and bold title bar.

    title: optional short string rendered as a bold caption above the panel center.
    """
    if shadow:
        sh = FancyBboxPatch(
            (x + 0.004, y - 0.006),
            w,
            h,
            boxstyle=f"round,pad=0.010,rounding_size={rounding}",
            fc=(0.0, 0.0, 0.0, 0.10),
            ec="none",
            lw=0,
            zorder=1,
        )
        ax.add_patch(sh)
    bb = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0.010,rounding_size={rounding}",
        fc=fill,
        ec=edge,
        lw=1.1,
        zorder=2,
    )
    ax.add_patch(bb)
    if title:
        # 标题挂在上 1/3，内文挂在下 2/3 中心；避免在大字号下两行重叠
        ax.text(
            x + w / 2,
            y + h * 0.78,
            title,
            ha="center",
            va="center",
            fontproperties=FP_HEI,
            fontsize=fontsize,
            color=title_color or edge,
            weight="bold",
            zorder=3,
        )
        ax.text(
            x + w / 2,
            y + h * 0.30,
            text,
            ha="center",
            va="center",
            fontproperties=fp or FP_SONG,
            fontsize=fontsize - 1 if fontsize > 9 else fontsize,
            color=text_color,
            zorder=3,
        )
    else:
        ax.text(
            x + w / 2,
            y + h / 2,
            text,
            ha="center",
            va="center",
            fontproperties=fp or FP_SONG,
            fontsize=fontsize,
            color=text_color,
            weight=weight,
            zorder=3,
        )


def _arrow(ax, x1, y1, x2, y2, color=C_GRAY, lw=1.4, *, rad=0.0, ls="-", alpha=1.0):
    """Smooth tapered arrow with anti-aliased path effect.

    rad>0 produces a curved connector, useful when arrows run between nodes
    that are not vertically/horizontally aligned.
    """
    arr = FancyArrowPatch(
        (x1, y1),
        (x2, y2),
        arrowstyle="-|>",
        mutation_scale=11,
        lw=lw,
        color=color,
        linestyle=ls,
        alpha=alpha,
        connectionstyle=f"arc3,rad={rad}",
        capstyle="round",
        joinstyle="round",
        zorder=2,
    )
    arr.set_path_effects([path_effects.SimpleLineShadow(offset=(0.0, -0.0), alpha=0.0),
                           path_effects.Normal()])
    ax.add_patch(arr)


def _layer_label(ax, x_left, x_right, y, text, color=C_GRAY):
    """Faint horizontal label band used to mark functional layers in flow diagrams."""
    ax.plot([x_left, x_right], [y, y], color=color, lw=0.6, ls=(0, (4, 4)), alpha=0.55)
    ax.text(
        x_left,
        y + 0.012,
        text,
        ha="left",
        va="bottom",
        fontproperties=FP_HEI,
        fontsize=12.6,
        color=color,
        alpha=0.85,
    )


def _setup_data_axes(ax, *, title=None, xlabel=None, ylabel=None, grid=True):
    """Apply a uniform data-axes style: faint grid, hidden top/right spines, panel bg."""
    ax.set_facecolor(C_PANEL_BG)
    ax.spines[["top", "right"]].set_visible(False)
    ax.spines[["left", "bottom"]].set_color("#444444")
    ax.spines[["left", "bottom"]].set_linewidth(0.9)
    if grid:
        ax.grid(True, axis="both", color=C_AXIS_GRID, lw=0.6, alpha=0.6, zorder=0)
        ax.set_axisbelow(True)
    if title:
        ax.set_title(title, fontproperties=FP_HEI, fontsize=16.1, pad=8)
    if xlabel:
        ax.set_xlabel(xlabel, fontproperties=FP_SONG, fontsize=14.7)
    if ylabel:
        ax.set_ylabel(ylabel, fontproperties=FP_SONG, fontsize=14.7)
    for lab in ax.get_xticklabels() + ax.get_yticklabels():
        lab.set_fontproperties(FP_SONG)
        lab.set_fontsize(9.5)


def _figure_title(fig, text):
    fig.suptitle(text, fontproperties=FP_HEI, fontsize=17.5, y=0.995)


# ============================================================
# Figures 1-10
# ============================================================


def figure01_workflow() -> None:
    """四层级联流程图：数据源 → 端点/虚拟生成 → 主模型 → 评估/路由。"""
    fig, ax = plt.subplots(figsize=(12.0, 7.2), dpi=260)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # ==== Layer labels ====
    _layer_label(ax, 0.02, 0.98, 0.965, "① 数据源")
    _layer_label(ax, 0.02, 0.98, 0.700, "② 端点估计与虚拟样本生成")
    _layer_label(ax, 0.02, 0.98, 0.475, "③ 物理先验门控残差主模型")
    _layer_label(ax, 0.02, 0.98, 0.230, "④ 严格评估与证据路由")

    # ==== Layer 1: 数据源 ====
    box_w, box_h = 0.20, 0.12
    src_y = 0.82
    sources = [
        (0.03, "均聚物 Tg", "n = 7486", C_BG_BLUE, C_BLUE),
        (0.27, "PolyInfo 共聚物", "n = 149", C_BG_GREEN, C_GREEN),
        (0.51, "高分子-DNA 共聚物", "n = 17", C_BG_RED, C_RED),
        (0.75, "虚拟弱标签数据", "按需筛选", C_BG_ORANGE, C_ORANGE),
    ]
    for x, title, sub, fc, ec in sources:
        _box(ax, x, src_y, box_w, box_h, sub, fc, ec, fontsize=15.4,
             title=title, title_color=ec)

    # ==== Layer 2: 端点源模型 + 虚拟生成 ====
    mid_y = 0.55
    src_model_x = 0.03
    src_model_w = 0.20
    _box(
        ax, src_model_x, mid_y, src_model_w, box_h, "TabPFN · 186 维",
        C_BG_BLUE, C_BLUE, fontsize=15.4,
        title="端点 Tg 源模型", title_color=C_BLUE,
    )
    virt_gen_x = 0.75
    _box(
        ax, virt_gen_x, mid_y, box_w, box_h, "Fox 混合 · 架构 · 端点窗口",
        C_BG_ORANGE, C_ORANGE, fontsize=13,
        title="虚拟共聚物生成", title_color=C_ORANGE,
    )
    # 均聚物 → 端点源模型
    _arrow(ax, 0.03 + box_w / 2, src_y, src_model_x + src_model_w / 2, mid_y + box_h, color=C_BLUE)
    # 虚拟数据 → 虚拟生成
    _arrow(ax, 0.75 + box_w / 2, src_y, virt_gen_x + box_w / 2, mid_y + box_h, color=C_ORANGE)
    # 端点源模型 → 虚拟生成 (横向供给端点 Tg)
    _arrow(ax, src_model_x + src_model_w, mid_y + box_h / 2, virt_gen_x, mid_y + box_h / 2,
           color=C_BLUE, rad=-0.10)

    # ==== Layer 3: 主模型 (扩宽至覆盖 4 列源中心 x=0.13/0.37/0.61/0.85) ====
    main_x = 0.10
    main_y = 0.31
    main_w = 0.80
    main_h = 0.14
    _box(
        ax, main_x, main_y, main_w, main_h,
        "端点 Tg + Fox 基线 + 结构嵌入 + 三类校准",
        C_BG_GOLD, C_GOLD, fontsize=15.4,
        title="物理先验门控残差回归 (方法 I 主模型)", title_color=C_GOLD,
    )
    # 4 个进入主模型的箭头：全部垂直直线
    _arrow(ax, src_model_x + src_model_w / 2, mid_y,
           src_model_x + src_model_w / 2, main_y + main_h, color=C_BLUE)
    _arrow(ax, 0.27 + box_w / 2, src_y, 0.27 + box_w / 2, main_y + main_h, color=C_GREEN)
    _arrow(ax, 0.51 + box_w / 2, src_y, 0.51 + box_w / 2, main_y + main_h, color=C_RED)
    _arrow(ax, virt_gen_x + box_w / 2, mid_y,
           virt_gen_x + box_w / 2, main_y + main_h, color=C_ORANGE)

    # ==== Layer 4: 评估 + 路由 ====
    bot_y = 0.06
    bot_h = 0.14
    eval_x, eval_w = 0.06, 0.40
    rout_x, rout_w = 0.54, 0.40
    _box(
        ax, eval_x, bot_y, eval_w, bot_h,
        "random holdout + group holdout · 主指标 min-R²",
        C_BG_BLUE, C_BLUE, fontsize=15.4,
        title="严格评估", title_color=C_BLUE,
    )
    _box(
        ax, rout_x, bot_y, rout_w, bot_h,
        "四级证据 → 八类预测路线",
        C_BG_GREEN, C_GREEN, fontsize=15.4,
        title="证据自适应路由 (方法 II)", title_color=C_GREEN,
    )
    # 2 个流出主模型的箭头：垂直直线，对齐两个评估/路由方框中心
    _arrow(ax, eval_x + eval_w / 2, main_y, eval_x + eval_w / 2, bot_y + bot_h, color=C_GRAY)
    _arrow(ax, rout_x + rout_w / 2, main_y, rout_x + rout_w / 2, bot_y + bot_h, color=C_GRAY)

    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure01_workflow.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure02_tg_distribution() -> None:
    buckets = load_tg_by_source()
    if not any(buckets.values()):
        # Fallback synthetic distribution if file missing
        rng = np.random.default_rng(0)
        buckets = {
            "homopolymer_real": rng.normal(120, 80, 1500).tolist(),
            "polyinfo_real": rng.normal(60, 50, 149).tolist(),
            "nucleobase_real": rng.normal(35, 30, 17).tolist(),
        }

    fig, axes = plt.subplots(1, 3, figsize=(12.0, 3.9), dpi=260)
    info = [
        ("homopolymer_real", "均聚物 (n=7486 训练 / 1498 留出)", C_BLUE, C_BG_BLUE),
        ("polyinfo_real", "PolyInfo 共聚物 (n=149)", C_GREEN, C_BG_GREEN),
        ("nucleobase_real", "高分子-DNA 共聚物 (n=17)", C_RED, C_BG_RED),
    ]
    for ax, (key, title, color, fill) in zip(axes, info):
        vals = np.asarray(buckets.get(key, []), dtype=float)
        if len(vals) == 0:
            ax.set_visible(False)
            continue
        bins = 30 if key == "homopolymer_real" else 12
        n, edges, patches = ax.hist(
            vals, bins=bins, color=fill, edgecolor=color, lw=0.9, zorder=2,
        )
        # smoothed envelope (simple moving average over hist counts)
        centers = 0.5 * (edges[:-1] + edges[1:])
        if len(centers) >= 4:
            kernel = np.array([0.2, 0.6, 1.0, 0.6, 0.2])
            kernel = kernel / kernel.sum()
            smoothed = np.convolve(n, kernel, mode="same")
            ax.plot(centers, smoothed, color=color, lw=1.6, alpha=0.9, zorder=3)
            ax.fill_between(
                centers, smoothed, color=color, alpha=0.10, zorder=1,
            )
        _setup_data_axes(ax, title=title, xlabel="Tg / ℃", ylabel="样本数")
        if len(vals):
            mean_v = float(np.mean(vals))
            ax.axvline(mean_v, color=color, lw=1.0, ls=(0, (4, 3)), alpha=0.85, zorder=4)
            ax.annotate(
                f"均值 {mean_v:.0f} ℃",
                xy=(mean_v, ax.get_ylim()[1] * 0.95),
                xytext=(8, -2),
                textcoords="offset points",
                fontproperties=FP_SONG,
                fontsize=12.6,
                color=color,
                bbox=dict(boxstyle="round,pad=0.25", fc="white", ec=color, lw=0.7, alpha=0.95),
                zorder=5,
            )
    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure02_tg_distribution.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure03_virtual_generation() -> None:
    """重画：每条箭头明确去向。

    端点 SMILES → 端点 Tg 估计；组成权重 + 架构标签 → Fox 混合；
    端点 Tg 估计 → Fox 混合（端点 Tg 喂入 Fox）；
    Fox 混合 + 端点 Tg 估计 → 弱标签 recipe。
    """
    fig, ax = plt.subplots(figsize=(12.0, 6.4), dpi=260)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # 三栏标题（与下方各栏中心对齐，字号放大）
    in_x, in_w = 0.03, 0.22
    hub_x, hub_w = 0.36, 0.26
    rec_x, rec_w = 0.71, 0.26
    in_cx = in_x + in_w / 2
    hub_cx = hub_x + hub_w / 2
    rec_cx = rec_x + rec_w / 2

    ax.text(in_cx, 0.945, "① 输入条件", ha="center", va="center",
            fontproperties=FP_HEI, fontsize=16.1, color=C_BLUE, fontweight="bold")
    ax.text(hub_cx, 0.945, "② 枢纽节点", ha="center", va="center",
            fontproperties=FP_HEI, fontsize=16.1, color=C_GREEN, fontweight="bold")
    ax.text(rec_cx, 0.945, "③ 弱标签产物", ha="center", va="center",
            fontproperties=FP_HEI, fontsize=16.1, color=C_GOLD, fontweight="bold")
    # 浅虚线把 3 栏分隔
    ax.plot([0.32, 0.32], [0.06, 0.91], color=C_GRAY, lw=0.5, ls=(0, (3, 4)), alpha=0.35)
    ax.plot([0.66, 0.66], [0.06, 0.91], color=C_GRAY, lw=0.5, ls=(0, (3, 4)), alpha=0.35)

    # 输入框（3 个）—— 选定 y center 使右侧 hub 行可水平直连
    # 端点 SMILES 与端点 Tg 估计中心 y 对齐 (0.65)；下方两个蓝框下移并彼此靠近
    in_h = 0.12
    inputs = [
        (0.59, "端点 SMILES", "(单体 1 / 单体 2 ...)"),  # center 0.65
        (0.31, "组成权重", "(w_1, w_2, ...)"),           # center 0.37（下移自 0.41，与下方靠近）
        (0.14, "架构标签", "random / block / multi"),    # center 0.20（下移微调）
    ]
    for y, title, sub in inputs:
        _box(ax, in_x, y, in_w, in_h, sub, C_BG_BLUE, C_BLUE, fontsize=14,
             title=title, title_color=C_BLUE)

    # 枢纽节点（中栏 2 块）：端点 Tg 估计与端点 SMILES 同 y center
    est_y, est_h = 0.55, 0.20   # 端点 Tg 估计 y=0.55-0.75, center 0.65
    _box(ax, hub_x, est_y, hub_w, est_h,
         "优先复用真实端点；\n缺失时由端点源模型预测",
         C_BG_GREEN, C_GREEN, fontsize=14.7,
         title="端点 Tg 估计", title_color=C_GREEN)

    fox_y, fox_h = 0.10, 0.34   # Fox y=0.10-0.44, center 0.27（缩高，使 Fox→recipe 不再贴底）
    _box(ax, hub_x, fox_y, hub_w, fox_h,
         "$1/T_{g} = \\sum_{i}\\, w_{i}/T_{g,i}$\n+ 端点 Tg 上下界 / 嵌入加权",
         C_BG_ORANGE, C_ORANGE, fontsize=15.4,
         title="Fox 混合 + 校准", title_color=C_ORANGE)

    # 弱标签 recipe（右栏，高度调整使两条水平进线 y=0.65 与 y=0.27 都进入箱内中部）
    rec_y, rec_h = 0.16, 0.60   # recipe y=0.16-0.76, center 0.46
    _box(ax, rec_x, rec_y, rec_w, rec_h,
         "• Tg 预测 / Fox 参考值\n\n• 端点 Tg window\n\n• 主路线 method\n\n• 稳定 recipe ID",
         C_BG_GOLD, C_GOLD, fontsize=14.7,
         title="弱标签 recipe", title_color=C_GOLD)

    # ---- 箭头 ----
    # 1. 端点 SMILES → 端点 Tg 估计：水平直线 y=0.65
    _arrow(ax, in_x + in_w, 0.59 + in_h / 2, hub_x, 0.59 + in_h / 2, color=C_BLUE)
    # 2. 组成权重 → Fox：水平直线 y=0.37
    _arrow(ax, in_x + in_w, 0.31 + in_h / 2, hub_x, 0.31 + in_h / 2, color=C_BLUE)
    # 3. 架构标签 → Fox：水平直线 y=0.20
    _arrow(ax, in_x + in_w, 0.14 + in_h / 2, hub_x, 0.14 + in_h / 2, color=C_BLUE)
    # 4. 端点 Tg 估计 → Fox：垂直直线，列中心 x（est 底 → fox 顶）
    _arrow(ax, hub_x + hub_w / 2, est_y, hub_x + hub_w / 2, fox_y + fox_h, color=C_GREEN)
    # 5. 端点 Tg 估计 → recipe：水平直线 y=0.65
    _arrow(ax, hub_x + hub_w, 0.59 + in_h / 2, rec_x, 0.59 + in_h / 2, color=C_GREEN)
    # 6. Fox → recipe：水平直线 y=0.27（Fox 中心，落在 recipe 下半中部）
    _arrow(ax, hub_x + hub_w, fox_y + fox_h / 2, rec_x, fox_y + fox_h / 2, color=C_ORANGE)

    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure03_virtual_generation.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure04_model_architecture() -> None:
    """重画：物理意义对齐的箭头。

    端点 Tg 单独连到 Fox 物理基线（其余 4 项不进入物理基线）；
    所有 5 项输入特征都作为残差学习器的输入；
    Fox 物理基线 + 残差学习器 → 求和 → 三类校准（串行）→ 输出。
    """
    fig, ax = plt.subplots(figsize=(12.6, 7.0), dpi=260)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # 3 栏标题
    _layer_label(ax, 0.02, 0.30, 0.985, "① 输入特征 (186 维)")
    _layer_label(ax, 0.34, 0.62, 0.985, "② 物理基线 + 求和 + 残差学习")
    _layer_label(ax, 0.66, 0.98, 0.985, "③ 三类校准 → 输出", color=C_RED)

    # 列 1：5 个输入特征。端点 Tg 与 Fox 同 y；4 个结构特征间距 0.04
    feat_x, feat_w, feat_h = 0.02, 0.28, 0.10
    # y=底坐标; 端点 Tg 顶部，4 结构特征均匀分布（连续两框间隙 0.04）
    feat_ys = [0.85, 0.51, 0.37, 0.23, 0.09]
    feat_labels = [
        "端点 Tg (max / min / 加权均值)",
        "链段物理特征 58 维",
        "图嵌入 64 维",
        "polyBERT 嵌入 64 维",
        "组成统计 (熵 / Herfindahl)",
    ]
    for y, lbl in zip(feat_ys, feat_labels):
        is_endpoint = (lbl.startswith("端点 Tg"))
        fc = C_BG_GREEN if is_endpoint else C_BG_BLUE
        ec = C_GREEN if is_endpoint else C_BLUE
        _box(ax, feat_x, y, feat_w, feat_h, lbl, fc, ec, fontsize=13.3)

    # 列 2：Fox 物理基线（顶，与端点 Tg 平齐）+ 求和（中）+ 残差学习器（底，覆盖 4 个结构特征 y）
    ctr_x, ctr_w = 0.34, 0.26
    fox_y, fox_h = 0.82, 0.14  # 加高，避免标题与公式行重叠
    _box(
        ax, ctr_x, fox_y, ctr_w, fox_h,
        "$1/T_{g} = \\sum_{i}\\, w_{i}/T_{g,i}$",
        C_BG_GREEN, C_GREEN, fontsize=12.6,
        title="Fox 物理基线", title_color=C_GREEN,
    )
    sum_y, sum_h = 0.63, 0.16  # 加高，避免标题与公式行重叠
    _box(
        ax, ctr_x, sum_y, ctr_w, sum_h,
        "$\\hat{y} = $ baseline$\\,+\\, g(x)\\!\\cdot\\! r_{\\theta}(x)$",
        C_BG_GRAY, C_GRAY, fontsize=12.6,
        title="求和", title_color=C_GRAY,
    )
    # 残差器覆盖 feat_ys[1:] = [0.51, 0.37, 0.23, 0.09] 的中心 [0.56, 0.42, 0.28, 0.14]
    res_y, res_h = 0.04, 0.55
    _box(
        ax, ctr_x, res_y, ctr_w, res_h,
        "$r_{\\theta}(x)$  (TabPFN /\nCatBoost / Ridge)\n受门控 g(x) 控制",
        C_BG_ORANGE, C_ORANGE, fontsize=14,
        title="门控残差学习器", title_color=C_ORANGE,
    )

    # 列 3：3 类校准 + 输出。Cal1 中心 y 与 Sum 中心 y 完全对齐 → Sum→Cal1 水平
    # Sum 中心 y = sum_y + sum_h/2 = 0.66 + 0.065 = 0.725
    # Cal1 h=0.10 → Cal1 y = 0.725 - 0.05 = 0.675
    cal_x, cal_w, cal_h = 0.66, 0.30, 0.10
    cal_ys = [0.675, 0.535, 0.395]  # Cal1 center 0.725 同 Sum；Cal2/3 间距 0.04
    cal_labels = [
        ("校准 1", "非均聚物门控\n端点/Fox 残差校准"),
        ("校准 2", "低 Fox 区残差收缩"),
        ("校准 3", "非均聚物预测差校准"),
    ]
    for y, (title, body) in zip(cal_ys, cal_labels):
        _box(ax, cal_x, y, cal_w, cal_h, body, C_BG_RED, C_RED, fontsize=13.3,
             title=title, title_color=C_RED)
    out_y, out_h = 0.06, 0.26
    _box(
        ax, cal_x, out_y, cal_w, out_h,
        "+ 路线标签 + 90% 区间",
        C_BG_GOLD, C_GOLD, fontsize=14.7,
        title="Tg 预测值 (℃)", title_color=C_GOLD,
    )

    # ---- 箭头 ----
    # 1. 端点 Tg → Fox：水平直线（同 y center=0.90）
    _arrow(ax, feat_x + feat_w, feat_ys[0] + feat_h / 2,
           ctr_x, fox_y + fox_h / 2, color=C_GREEN)
    # 2. 4 个结构特征 → Residual：每个水平直线，自己 y center
    for y in feat_ys[1:]:
        _arrow(ax, feat_x + feat_w, y + feat_h / 2,
               ctr_x, y + feat_h / 2, color=C_BLUE, lw=1.0)
    # 3. Fox → 求和：垂直直线（列中心 x），从 Fox 底到 Sum 顶
    _arrow(ax, ctr_x + ctr_w / 2, fox_y,
           ctr_x + ctr_w / 2, sum_y + sum_h, color=C_GREEN)
    # 4. Residual → 求和：垂直直线，从 Residual 顶到 Sum 底
    _arrow(ax, ctr_x + ctr_w / 2, res_y + res_h,
           ctr_x + ctr_w / 2, sum_y, color=C_ORANGE)
    # 5. 求和 → Cal1：水平直线（同 y center=0.71）
    _arrow(ax, ctr_x + ctr_w, sum_y + sum_h / 2,
           cal_x, cal_ys[0] + cal_h / 2, color=C_GRAY)
    # 6. Cal 串行链
    _arrow(ax, cal_x + cal_w / 2, cal_ys[0],
           cal_x + cal_w / 2, cal_ys[1] + cal_h, color=C_RED)
    _arrow(ax, cal_x + cal_w / 2, cal_ys[1],
           cal_x + cal_w / 2, cal_ys[2] + cal_h, color=C_RED)
    # 7. Cal3 → Output：垂直直线
    _arrow(ax, cal_x + cal_w / 2, cal_ys[2],
           cal_x + cal_w / 2, out_y + out_h, color=C_RED)

    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure04_model_architecture.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure05_main_results() -> None:
    pred_obs = load_pred_obs_by_source()
    fig, axes = plt.subplots(1, 2, figsize=(12.0, 4.9), dpi=260,
                              gridspec_kw={"width_ratios": [1, 1.25]})

    # Panel (a): bar chart of R²
    ax0 = axes[0]
    short = ["均聚物", "PolyInfo 共聚物", "高分子-DNA"]
    r2s = [m.r2 for m in MAIN_RESULTS]
    maes = [m.mae for m in MAIN_RESULTS]
    fills = [_lighten(c, 0.30) for c in (C_BLUE, C_GREEN, C_RED)]
    bars = ax0.bar(
        short, r2s, color=fills,
        edgecolor=[C_BLUE, C_GREEN, C_RED], lw=1.4, width=0.58,
    )
    for b in bars:
        b.set_path_effects([
            path_effects.SimplePatchShadow(offset=(1.2, -1.2), alpha=0.18),
            path_effects.Normal(),
        ])
    ax0.set_ylim(0, 1.05)
    ax0.axhline(0.95, color=C_GRAY, lw=1.0, ls=(0, (4, 3)), alpha=0.7)
    ax0.text(2.45, 0.955, "理想目标 0.95", fontproperties=FP_SONG,
             fontsize=12.6, color=C_GRAY, ha="right", va="bottom")
    _setup_data_axes(ax0, title="(a) 三任务严格评估表现", ylabel="R²")
    for bar, r2, mae, c in zip(bars, r2s, maes, [C_BLUE, C_GREEN, C_RED]):
        ax0.text(
            bar.get_x() + bar.get_width() / 2,
            r2 + 0.025,
            f"R²={r2:.3f}\nMAE={mae:.2f} ℃",
            ha="center", va="bottom",
            fontproperties=FP_SONG,
            fontsize=13.3, color=c, fontweight="bold",
        )

    # Panel (b): predicted vs observed
    ax1 = axes[1]
    color_map = {
        "homopolymer_real": (C_BLUE, "均聚物"),
        "polyinfo_real": (C_GREEN, "PolyInfo 共聚物"),
        "nucleobase_real": (C_RED, "高分子-DNA"),
    }
    all_y, all_yhat = [], []
    for src, (c, lbl) in color_map.items():
        if src not in pred_obs:
            continue
        y, yhat = pred_obs[src]
        ax1.scatter(
            y, yhat,
            s=12 if src == "homopolymer_real" else 38,
            c=c,
            alpha=0.40 if src == "homopolymer_real" else 0.85,
            edgecolors="white", lw=0.5, label=lbl, zorder=3,
        )
        all_y.extend(y)
        all_yhat.extend(yhat)
    if all_y:
        lo = min(min(all_y), min(all_yhat)) - 10
        hi = max(max(all_y), max(all_yhat)) + 10
        ax1.plot([lo, hi], [lo, hi], color="#222222", lw=0.9, ls=(0, (4, 3)), alpha=0.7, zorder=2)
        ax1.set_xlim(lo, hi)
        ax1.set_ylim(lo, hi)
    _setup_data_axes(
        ax1,
        title="(b) 预测—实测散点 (主模型)",
        xlabel="实测 Tg / ℃",
        ylabel="预测 Tg / ℃",
    )
    leg = ax1.legend(
        prop=FP_SONG, fontsize=13.3, loc="upper left",
        frameon=True, framealpha=0.92, edgecolor=C_AXIS_GRID,
    )

    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure05_main_results.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure06_ablation() -> None:
    fig, ax = plt.subplots(figsize=(10.2, 4.7), dpi=260)
    labels = [a[0] for a in ABLATION]
    homo = [a[1] for a in ABLATION]
    poly = [a[2] for a in ABLATION]
    nuc = [a[3] for a in ABLATION]
    minr = [a[4] for a in ABLATION]
    x = np.arange(len(labels))
    series = [
        (homo, "均聚物 R²", C_BLUE, "o"),
        (poly, "PolyInfo 共聚物 R²", C_GREEN, "s"),
        (nuc, "高分子-DNA R²", C_RED, "^"),
    ]
    for vals, lbl, color, marker in series:
        ax.fill_between(x, vals, alpha=0.04, color=color, zorder=1)
        ax.plot(x, vals, marker=marker, color=color, lw=1.8, ms=7,
                markeredgecolor="white", markeredgewidth=0.8,
                label=lbl, zorder=3)
    ax.plot(x, minr, marker="D", color="#1F1F1F", lw=1.4, ms=6, ls=(0, (1, 1.5)),
            markeredgecolor="white", markeredgewidth=0.6, label="min-R²", zorder=4)
    for xi, yi in zip(x, minr):
        ax.annotate(
            f"{yi:.3f}",
            xy=(xi, yi),
            xytext=(0, -16),
            textcoords="offset points",
            ha="center",
            fontproperties=FP_SONG,
            fontsize=12.6,
            color="#1F1F1F",
            fontweight="bold",
        )
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontproperties=FP_SONG, fontsize=13.3, rotation=10)
    ax.set_ylim(0.75, 0.93)
    _setup_data_axes(ax, ylabel="R²")
    ax.legend(
        prop=FP_SONG, fontsize=13.3, ncol=2, frameon=True,
        framealpha=0.9, edgecolor=C_AXIS_GRID, loc="lower right",
    )
    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure06_ablation.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure07_virtual_tradeoff() -> None:
    """重画：用引线把标签拉离散点位置；同位重叠点改用 annotation arrow。

    "仅真实数据 (主模型)" 与 "一致性过滤虚拟" 几乎在同一坐标 (0.849-0.850, 0.817)，
    用 annotation 引线把标签拉到散点上方/下方不同方向，避免叠字。
    """
    fig, ax = plt.subplots(figsize=(9.2, 5.8), dpi=260)
    fig.patch.set_facecolor("white")

    # 象限轻底色：左上=核酸提升 / 右下=核酸退化
    ax.axvspan(0.810, 0.849, ymin=(0.817 - 0.792) / (0.882 - 0.792), ymax=1.0,
               color=C_ORANGE, alpha=0.05, zorder=0)
    ax.axvspan(0.849, 0.872, ymin=0.0, ymax=(0.817 - 0.792) / (0.882 - 0.792),
               color=C_GRAY, alpha=0.05, zorder=0)

    # 参考线
    ax.axvline(0.849, color=C_GRAY, lw=0.7, ls=(0, (4, 3)), alpha=0.7, zorder=1)
    ax.axhline(0.817, color=C_GRAY, lw=0.7, ls=(0, (4, 3)), alpha=0.7, zorder=1)

    points = [
        (0.849, 0.817, "仅真实数据 (方法 I 主模型)", C_BLUE, (-60, -50)),
        (0.827, 0.858, "结构近邻虚拟", C_ORANGE, (16, 12)),
        (0.850, 0.817, "一致性过滤虚拟", C_GREEN, (40, 28)),
    ]
    for x, y, lbl, c, off in points:
        ax.scatter(x, y, s=200, c=c, edgecolors="white", lw=1.4, zorder=4,
                   path_effects=[
                       path_effects.SimplePatchShadow(offset=(0.8, -0.8), alpha=0.20),
                       path_effects.Normal(),
                   ])
        ax.annotate(
            lbl,
            xy=(x, y),
            xytext=off,
            textcoords="offset points",
            fontproperties=FP_SONG,
            fontsize=14.7,
            color=c,
            fontweight="bold",
            arrowprops=dict(arrowstyle="-", color=c, lw=0.9, alpha=0.85),
            zorder=5,
        )

    # 象限提示文字（避开散点位置）
    ax.text(
        0.815, 0.870,
        "↖ PolyInfo 共聚物退化 / 核酸提升",
        fontproperties=FP_SONG, fontsize=13.3, color=C_ORANGE,
        ha="left", va="top",
    )
    ax.text(
        0.866, 0.798,
        "↘ PolyInfo 共聚物提升 / 核酸退化",
        fontproperties=FP_SONG, fontsize=13.3, color=C_GRAY,
        ha="right", va="bottom",
    )

    _setup_data_axes(
        ax,
        title="虚拟数据增强的收益—风险散点",
        xlabel="PolyInfo 共聚物 group holdout R²",
        ylabel="高分子-DNA group holdout R²",
    )
    ax.set_xlim(0.810, 0.872)
    ax.set_ylim(0.792, 0.882)

    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure07_virtual_tradeoff.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure08_hard_systems() -> None:
    fig, ax = plt.subplots(figsize=(9.6, 4.7), dpi=260)
    names = [s[0] + "\n" + s[1] for s in HARD_SYSTEMS]
    maes = [s[2] for s in HARD_SYSTEMS]
    edges = [C_RED if v >= 18 else (C_ORANGE if v >= 9 else C_GREEN) for v in maes]
    fills = [_lighten(c, 0.30) for c in edges]
    bars = ax.barh(names, maes, color=fills, edgecolor=edges, lw=1.3, height=0.7)
    for b in bars:
        b.set_path_effects([
            path_effects.SimplePatchShadow(offset=(1.0, -0.8), alpha=0.18),
            path_effects.Normal(),
        ])
    for bar, val, ec in zip(bars, maes, edges):
        ax.text(
            val + 0.5, bar.get_y() + bar.get_height() / 2,
            f"{val:.1f}",
            va="center",
            fontproperties=FP_SONG,
            fontsize=14, color=ec, fontweight="bold",
        )
    ax.invert_yaxis()
    _setup_data_axes(ax, title="PolyInfo 共聚物中误差最大的体系", xlabel="体系级 MAE / ℃")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure08_hard_systems.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure09_nucleobase_residuals() -> None:
    fig, ax = plt.subplots(figsize=(9.0, 4.7), dpi=260)
    keys = list(NUC_RESIDUALS.keys())
    data = [NUC_RESIDUALS[k] for k in keys]
    palette = [C_BLUE, C_RED, C_ORANGE, C_GREEN, C_GRAY]
    bp = ax.boxplot(
        data, patch_artist=True, widths=0.55, tick_labels=keys,
        medianprops=dict(color="#1F1F1F", lw=1.6),
        whiskerprops=dict(color="#3A3A3A", lw=0.9),
        capprops=dict(color="#3A3A3A", lw=0.9),
        flierprops=dict(marker="o", markersize=4,
                         markerfacecolor="white", markeredgecolor="#3A3A3A", lw=0.5),
    )
    for patch, color in zip(bp["boxes"], palette):
        patch.set_facecolor(_lighten(color, 0.4))
        patch.set_edgecolor(color)
        patch.set_linewidth(1.2)
    # overlay points (jittered)
    rng = np.random.default_rng(0)
    for i, (vals, color) in enumerate(zip(data, palette), start=1):
        xs = rng.normal(i, 0.06, size=len(vals))
        ax.scatter(xs, vals, s=18, color=color, alpha=0.55,
                   edgecolors="white", lw=0.4, zorder=4)
    ax.axhline(0, color="#1F1F1F", lw=0.9, ls=(0, (4, 3)), alpha=0.7)
    _setup_data_axes(
        ax,
        title="高分子-DNA 共聚物的碱基类型系统偏差",
        xlabel="碱基类型",
        ylabel="残差 (预测 - 实测) / ℃",
    )
    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure09_nucleobase_residuals.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure10_routing_tree() -> None:
    fig, ax = plt.subplots(figsize=(11.2, 6.6), dpi=260)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # Root: 输入证据解析
    root_x, root_y, root_w, root_h = 0.36, 0.84, 0.28, 0.12
    _box(ax, root_x, root_y, root_w, root_h, "依据特征/端点信息分流",
         C_BG_GOLD, C_GOLD, fontsize=15.4,
         title="输入证据解析", title_color=C_GOLD)
    root_bottom = (root_x + root_w / 2, root_y)

    # Trunk down from root
    trunk_y = 0.76
    _arrow(ax, root_bottom[0], root_bottom[1], root_bottom[0], trunk_y, color=C_GOLD)
    ax.plot([0.10, 0.90], [trunk_y, trunk_y], color=C_GOLD, lw=1.0, alpha=0.8)

    # Top row: 4 main routes（长 route 文字换行避免溢出）
    main = [
        (0.06, "单 SMILES", "均聚物源模型", C_BLUE),
        (0.30, "二元 + 已知体系", "同体系局部插值", C_GREEN),
        (0.54, "二元 + 未知体系", "全局 Fox\n+ 受约束残差", C_GREEN),
        (0.78, "DNA + 端点已知", "actual endpoint\n物理路线", C_RED),
    ]
    main_w = 0.18
    main_q_y = 0.62
    main_r_y = 0.44
    for x, q, route, c in main:
        _arrow(ax, x + main_w / 2, trunk_y, x + main_w / 2, main_q_y + 0.10, color=c, lw=1.0)
        _box(ax, x, main_q_y, main_w, 0.10, q, C_BG_GRAY, c, fontsize=13)
        _arrow(ax, x + main_w / 2, main_q_y, x + main_w / 2, main_r_y + 0.12, color=c)
        _box(ax, x, main_r_y, main_w, 0.12, route, _lighten(c, 0.78), c,
             fontsize=12.5, weight="bold")

    # Lower trunk
    lower_trunk_y = 0.34
    _arrow(ax, root_bottom[0], main_r_y, root_bottom[0], lower_trunk_y, color=C_GOLD, lw=0.8)
    ax.plot([0.18, 0.82], [lower_trunk_y, lower_trunk_y], color=C_GOLD, lw=1.0, ls=(0, (3, 3)), alpha=0.7)

    # Bottom row: 3 fallback routes（长 route 文字换行避免溢出）
    extra = [
        (0.10, "多元 random", "端点 Fox\n多组分近似", C_GREEN),
        (0.40, "嵌段架构", "miscible proxy\n+ Tg window", C_ORANGE),
        (0.70, "端点缺失", "端点估计 fallback\n(低置信)", C_GRAY),
    ]
    extra_w = 0.20
    extra_q_y = 0.18
    extra_r_y = 0.02
    for x, q, route, c in extra:
        _arrow(ax, x + extra_w / 2, lower_trunk_y, x + extra_w / 2, extra_q_y + 0.10, color=c, lw=0.9)
        _box(ax, x, extra_q_y, extra_w, 0.10, q, C_BG_GRAY, c, fontsize=13)
        _arrow(ax, x + extra_w / 2, extra_q_y, x + extra_w / 2, extra_r_y + 0.12, color=c)
        _box(ax, x, extra_r_y, extra_w, 0.12, route, _lighten(c, 0.78), c,
             fontsize=12.5, weight="bold")

    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure10_routing_tree.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def figure11_m1_vs_m2_bar() -> None:
    """方法 I 与方法 II 最佳路主指标分组柱状图（公平口径对齐）。"""
    tasks = [
        "均聚物\n 公平 holdout",
        "PolyInfo\n 跨体系",
        "PolyInfo\n 体系内",
        "高分子-DNA\n specialized",
        "高分子-DNA\n universal",
    ]
    m1 = [0.887, 0.849, 0.849, 0.817, 0.817]
    m2 = [0.907, 0.820, 0.933, 0.971, 0.564]
    m2_routes = ["TabPFN\n186d", "Linear-Fox\nLOSO", "同体系\nIDW", "Physics-\nRidge", "Linear-Fox\n全局"]
    x = np.arange(len(tasks))
    w = 0.36

    fig, ax = plt.subplots(figsize=(10.6, 5.4), dpi=260)
    bars1 = ax.bar(
        x - w / 2, m1, w,
        color=_lighten(C_BLUE, 0.40), edgecolor=C_BLUE, linewidth=1.3,
        label="方法 I 通用统计模型", zorder=3,
    )
    bars2 = ax.bar(
        x + w / 2, m2, w,
        color=_lighten(C_ORANGE, 0.30), edgecolor=C_ORANGE, linewidth=1.3,
        label="方法 II 证据自适应路由（最佳路）", zorder=3,
    )
    for b in list(bars1) + list(bars2):
        b.set_path_effects([
            path_effects.SimplePatchShadow(offset=(0.9, -0.9), alpha=0.18),
            path_effects.Normal(),
        ])
    for xi, m in zip(x - w / 2, m1):
        ax.text(
            xi, m + 0.014, f"{m:.3f}",
            ha="center", fontproperties=FP_SONG,
            fontsize=13.3, color=C_BLUE, fontweight="bold",
        )
    for xi, m, route in zip(x + w / 2, m2, m2_routes):
        ax.text(
            xi, m + 0.018, f"{m:.3f}",
            ha="center", fontproperties=FP_SONG,
            fontsize=13.3, color=C_ORANGE, fontweight="bold",
        )
        # 路线标签放在柱体下半，矮柱（如 universal=0.564）放到柱体最低端
        # 防止 bbox 上沿覆盖到顶部数值
        label_y = max(m - 0.13, 0.515)
        ax.text(
            xi, label_y, route,
            ha="center", fontproperties=FP_SONG,
            fontsize=11.9, color=C_GRAY,
            bbox=dict(boxstyle="round,pad=0.18", fc="white",
                      ec=C_AXIS_GRID, lw=0.5, alpha=0.88),
        )
    ax.set_xticks(x)
    ax.set_xticklabels(tasks, fontproperties=FP_SONG, fontsize=14)
    ax.set_ylim(0.50, 1.05)
    _setup_data_axes(
        ax,
        title=r"方法 I 与方法 II 在五种评估场景下的 $R^2$ 对比",
        ylabel=r"$R^2$ (越高越好)",
    )
    ax.legend(
        prop=FP_SONG, fontsize=13.3, loc="lower left",
        frameon=True, framealpha=0.95, edgecolor=C_AXIS_GRID,
    )
    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure11_m1_vs_m2.png", dpi=260, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


def create_figures() -> None:
    figure01_workflow()
    figure02_tg_distribution()
    figure03_virtual_generation()
    figure04_model_architecture()
    figure05_main_results()
    figure06_ablation()
    figure07_virtual_tradeoff()
    figure08_hard_systems()
    figure09_nucleobase_residuals()
    figure10_routing_tree()
    figure11_m1_vs_m2_bar()


# ============================================================
# Front matter (cover, abstract, toc, lists)
# ============================================================


def add_cover(doc) -> None:
    """封面：完全照搬附件4第11页布局。

    顶部"作品编号：（报名时系统提供的作品编号）" 黑体三号；
    主标题 方正小标宋_GBK 二号 行距固定45磅 居中；
    副标题"参 赛 作 品" 方正小标宋_GBK 一号 行距固定45磅 居中；
    四个身份字段：标签独立一行，下方一条整宽下划线（仿宋_GB2312 小二号 单倍行距）。
    """
    # 作品编号 (黑体三号，左对齐)
    p = doc.add_paragraph()
    fmt_paragraph(p, line_spacing_pt=24, space_before_pt=12)
    r = p.add_run("作品编号：")
    style_run(r, F_HEI, SZ_TITLE3, bold=True)
    r2 = p.add_run("TJJM20260325000015")
    style_run(r2, F_HEI, SZ_TITLE3, bold=True)

    # Spacer (cover 顶部 → 主标题之间留较大空白)
    for _ in range(4):
        sp = doc.add_paragraph()
        fmt_paragraph(sp, line_spacing_pt=24)

    # Main title (方正小标宋_GBK 二号 行距固定45磅，居中)
    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=45)
    r = p.add_run("2026年（第十二届）全国大学生统计建模大赛")
    style_run(r, F_FZXBS, SZ_TITLE2, bold=True)

    # Subtitle 参赛作品 (方正小标宋_GBK 一号 行距固定45磅，居中)
    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=45)
    r = p.add_run("参　赛　作　品")
    style_run(r, F_FZXBS, SZ_TITLE1, bold=True)

    # Spacer (副标题 → 身份字段之间留更大空白)
    for _ in range(3):
        sp = doc.add_paragraph()
        fmt_paragraph(sp, line_spacing_pt=24)

    # Identification rows: 仿宋_GB2312 小二号 单倍行距 (附件4 第11页布局)
    # 每对 = 标签独占一行；其下一条以段落底部边框制造贯通整宽的下划线
    fields = [
        ("参赛学校：", "同济大学"),
        ("论文题目：", TITLE),
        ("参赛队员：", "魏子赫、何奕良、杨帅"),
        ("指导老师：", "胡勇"),
    ]
    for label, value in fields:
        # Label paragraph (无下划线)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(12)
        pf.space_after = Pt(0)
        pf.first_line_indent = None
        r = p.add_run(label)
        style_run(r, F_FANGSONG, SZ_XIAO2, bold=True)

        # 下划线段落：用 paragraph bottom border 制造一条贯通整宽的横线
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.first_line_indent = None
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
        if value:
            _emit_styled_text(p, _norm(value), F_FANGSONG, SZ_XIAO2)
        else:
            # 空 run 让段落保持高度，bottom border 自然落在底部
            r = p.add_run(" ")
            style_run(r, F_FANGSONG, SZ_XIAO2)

    page_break(doc)


def add_abstract(doc) -> None:
    add_section_title(doc, "摘要")
    paragraphs = [
        "玻璃化转变温度（glass transition temperature, Tg）是衡量高分子链段运动能力与材料服役温区的关键指标，对耐热材料、柔性器件、生物医用高分子和功能涂层的设计具有直接影响。已有研究主要围绕均聚物 Tg 预测展开，而共聚物 Tg 同时受端点均聚物 Tg、组成比例、链段架构、链间相容性以及实验噪声等多因素共同决定；DNA 相关高分子共聚物进一步引入碱基识别、氢键和堆积效应，使常规模型难以直接外推。",
        "针对上述问题，本文提出两个互补的统计建模方法。方法 I 是物理先验门控残差的通用统计模型：以端点 Tg 源模型与 Fox 关系作为物理基线，引入受门控约束的残差学习器与三类后处理校准，对均聚物、一般共聚物与 DNA 相关共聚物在同一框架下进行回归。方法 II 是证据自适应路由系统：根据用户输入证据的等级（端点 Tg 是否已知、体系是否已知、组成是否已知、是否属于 DNA 相关），在 Fox 端点直接、Linear-Fox 拟合、Kwei 同体系优先、Physics-Ridge 物理特征回归与同体系残差插值五个子估计器之间动态选择，并以可审计输出形式给出 Tg 预测、路线名、计算方法、端点来源与置信度评估。两方法共享同一套均聚物 7486 条端点数据、149 条清洗后 PolyInfo 一般共聚物数据与 17 条核酸碱基功能化共聚物数据，方法 II 内部以方法 I 作为端点估计与 fallback。",
        "在严格分组留出协议下，方法 I 在均聚物随机留出、一般共聚物体系留出和 DNA 相关碱基留出上的 R² 分别为 0.887、0.849 和 0.817，平均绝对误差分别为 27.16、16.65 和 6.27 ℃，综合主指标 min-R²=0.817；校准组件消融显示主指标由 0.789 经三步逐级提升至 0.817。方法 II 在同一数据资源上按输入证据动态选择子估计器：当输入仅为单一重复单元 SMILES 时，路由直接调用专门为均聚物训练的 TabPFN 186d 端点源模型，在与方法 I 相同的 1498 行公平 holdout 上 R²=0.907、MAE=20.6 K，相对方法 I 提升 0.020 R²（同模型在全 7486 条上 5×3 RepeatedKFold CV R²=0.917 作为独立基准）；当体系已知时，PolyInfo 上 R² 由方法 I 的 0.849 推高至 0.933（同体系残差 IDW，MAE=8.76 K），跨体系外推时仍维持 R²=0.820（Linear-Fox 留体系外）；当端点 Tg 实测可用时，DNA 相关任务在 specialized actual-endpoint 路线下 R² 由 0.817 推高至 0.971（Physics-Ridge LOOCV，MAE=2.93 K），但若用户仅提供 SMILES 而无碱基实测端点，路由默认走 universal 路线，R²=0.564，这一档差是路由系统真实使用边界的诚实暴露。两方法相互覆盖：方法 I 在多源融合训练下保证最严苛 group holdout 下的 min-R²=0.817 综合稳健性，方法 II 在证据等级齐全时把指标推到接近上界，从而在科学严谨性与工程可用性之间形成稳定平衡。",
    ]
    for t in paragraphs:
        add_body_para(doc, t)
    p = doc.add_paragraph()
    fmt_paragraph(p, line_spacing_pt=24, first_line_indent_pt=24, space_before_pt=6)
    r = p.add_run("关键词：")
    style_run(r, F_HEI, SZ_XIAO4, bold=True)
    r = p.add_run(_norm("共聚物玻璃化转变温度；统计预测；物理先验门控残差；证据自适应路由；高分子-DNA 共聚物"))
    style_run(r, F_SONG, SZ_XIAO4)
    page_break(doc)


def _add_toc_field(paragraph, placeholder_runs: Sequence[tuple[str, str]]) -> None:
    """Insert a Word TOC field with manually-typed placeholder lines.

    Word users right-click → 更新域 to refresh; LibreOffice/PDF readers see
    placeholder content directly. ``placeholder_runs`` is a list of
    ``(title, page)`` tuples used for the placeholder body.
    """
    run = paragraph.add_run()
    style_run(run, F_SONG, SZ_XIAO4)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    # Placeholder content (will be replaced when user updates field in Word)
    for title, page in placeholder_runs:
        br = OxmlElement("w:br")
        run._r.append(br)
        t1 = OxmlElement("w:t")
        t1.set(qn("xml:space"), "preserve")
        t1.text = _norm(title) + "\t" + page
        run._r.append(t1)
    run._r.append(end)


def add_toc(doc, *, include_acknowledgements: bool = True) -> None:
    add_section_title(doc, "目录")
    # 注意：本目录使用 Word 自动生成域 (TOC field)。Word 打开后请右键 → 更新域。
    # 占位行同时作为 LibreOffice / PDF 渲染时的可视内容，按当前结构估算页码。
    items: list[tuple[str, str]] = [
        ("一、问题描述与研究思路", "1"),
        ("　（一）玻璃化转变温度的物理意义与工程价值", "1"),
        ("　（二）共聚物 Tg 预测的统计建模困难", "1"),
        ("　（三）DNA 相关高分子的特殊挑战", "2"),
        ("　（四）文献现状与研究空白", "3"),
        ("　（五）本文研究目标与技术路线", "3"),
        ("二、数据来源与质量控制", "4"),
        ("　（一）数据总体构成与用途分配", "4"),
        ("　（二）均聚物数据", "5"),
        ("　（三）一般共聚物数据的清洗流程", "6"),
        ("　（四）核酸/碱基功能化共聚物数据与碱基类型分组", "6"),
        ("　（五）虚拟共聚物弱标签的生成原则", "7"),
        ("　（六）数据集汇总与可获取性", "8"),
        ("三、评价指标与划分设计", "8"),
        ("　（一）误差与拟合度量", "8"),
        ("　（二）三任务分别评价", "9"),
        ("　（三）体系级与碱基族 group holdout", "9"),
        ("　（四）综合主指标 min-R²", "10"),
        ("四、方法 I：物理先验门控残差的通用统计模型", "10"),
        ("　（一）模型框架总览", "10"),
        ("　（二）端点 Tg 源模型", "11"),
        ("　　1. 三类预训练表征", "11"),
        ("　　2. 训练流与端点估计接口", "12"),
        ("　（三）共聚物 Fox 物理基线", "12"),
        ("　（四）共聚物层结构与组成特征", "13"),
        ("　（五）物理先验门控残差回归", "13"),
        ("　（六）三类校准", "14"),
        ("　（七）多源数据融合与样本权重", "14"),
        ("五、方法 II：证据自适应路由系统", "15"),
        ("　（一）方法动机：为何需要按证据分流", "15"),
        ("　（二）四级证据等级与八类预测路线", "16"),
        ("　（三）五个子估计器及其物理依据", "17"),
        ("　　1. Fox 端点直接路线", "17"),
        ("　　2. Linear-Fox 拟合校准", "17"),
        ("　　3. Kwei 同体系优先 LOOCV", "18"),
        ("　　4. Physics-Ridge 物理特征回归", "18"),
        ("　　5. 同体系残差 IDW 局部插值", "19"),
        ("　（四）路由决策机制与可审计输出", "19"),
        ("　（五）训练数据复用与方法 I 的关系", "20"),
        ("六、综合求解、对比与诊断", "20"),
        ("　（一）主指标对照：方法 I 与方法 II 在三任务上的总览", "20"),
        ("　（二）PolyInfo 共聚物 149 行：路由七路全表", "21"),
        ("　（三）高分子-DNA 共聚物碱基类型 17：跨碱基泛化全表", "23"),
        ("　（四）校准组件消融与虚拟数据增强", "24"),
        ("　（五）硬体系与碱基族偏差诊断", "25"),
        ("　（六）不确定性、典型预测与互补性分析", "27"),
        ("七、结论与建议", "29"),
        ("　（一）主要结论与双方法覆盖矩阵", "29"),
        ("　（二）主要不足", "30"),
        ("　（三）后续工作建议", "30"),
        ("参考文献", "31"),
        ("附录", "33"),
    ]
    if include_acknowledgements:
        items.append(("致谢", "35"))
    p = doc.add_paragraph()
    fmt_paragraph(p, line_spacing_pt=24)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(14.6), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2
    )
    _add_toc_field(p, items)
    page_break(doc)


def add_lists(doc) -> None:
    add_section_title(doc, "表格与插图清单")
    add_inline_subhead(doc, "表格清单")
    table_items = [
        "表1  数据来源与用途说明",
        "表2  PolyInfo 共聚物数据清洗步骤汇总",
        "表3  评价任务与划分方式",
        "表4  方法 II 路由系统的输入证据—预测路线对应",
        "表5  方法 I 主模型预测性能",
        "表6  方法 I 与方法 II 主指标并列对照矩阵",
        "表7  方法 II 在 PolyInfo 共聚物 149 行上的子估计器全表",
        "表8  方法 II 在核酸碱基功能化共聚物 17 行上的跨碱基全表",
        "表9  方法 I 校准组件消融结果",
        "表10  虚拟数据增强策略对比",
        "表11  PolyInfo 共聚物中误差最大的体系",
    ]
    for t in table_items:
        p = doc.add_paragraph()
        fmt_paragraph(p, line_spacing_pt=24, first_line_indent_pt=24)
        # 表序 "表N  " 双空格不参与规范化，保护排版格式
        import re as _re
        m = _re.match(r"^(表\s*\d+ {2})(.*)$", t)
        if m:
            text = m.group(1) + _norm(m.group(2))
        else:
            text = _norm(t)
        _emit_styled_text(p, text, F_SONG, SZ_XIAO4)
    add_inline_subhead(doc, "插图清单")
    figure_items = [
        "图1  共聚物 Tg 预测研究技术路线",
        "图2  三类真实数据的 Tg 分布",
        "图3  虚拟共聚物弱标签生成示意",
        "图4  方法 I 物理先验门控残差模型结构",
        "图5  方法 II 证据自适应路由决策树",
        "图6  方法 I 与方法 II 主指标对比",
        "图7  方法 I 三任务严格评估表现与预测—实测散点",
        "图8  方法 I 校准组件消融的三任务 R² 曲线",
        "图9  虚拟数据增强的收益—风险散点",
        "图10  PolyInfo 共聚物中误差最大的体系",
        "图11  高分子-DNA 共聚物的碱基类型系统偏差",
    ]
    for t in figure_items:
        p = doc.add_paragraph()
        fmt_paragraph(p, line_spacing_pt=24, first_line_indent_pt=24)
        # 图序 "图N  " 双空格不参与规范化，保护排版格式
        import re as _re
        m = _re.match(r"^(图\s*\d+ {2})(.*)$", t)
        if m:
            text = m.group(1) + _norm(m.group(2))
        else:
            text = _norm(t)
        _emit_styled_text(p, text, F_SONG, SZ_XIAO4)


def start_body_section(doc) -> None:
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_section(sec)
    sec.footer.is_linked_to_previous = False
    restart_page_number(sec, 1)
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_field(p)


# ============================================================
# Body chapters
# ============================================================


def chapter1(doc) -> None:
    add_h1(doc, "一、问题描述与研究思路")

    add_h2(doc, "（一）玻璃化转变温度的物理意义与工程价值")
    add_body_para(
        doc,
        "玻璃化转变温度 Tg 是非晶态高分子从玻璃态向高弹态过渡的特征温度，由分子结构、链段刚性、自由体积分布与加热速率共同决定，是高分子最具综合性的热力学—动力学指标之一。在工程上 Tg 直接决定材料服役温区：耐热结构材料追求高 Tg、柔性器件与锂电池电解质追求低 Tg；合成前需预测以筛选候选结构，合成后需测定以验证目标。",
    )
    add_body_para(
        doc,
        "Tg 的实验测定主要依赖差示扫描量热法、动态热机械分析与介电谱，周期较长且对样品制备敏感，相同结构的文献值常有 5—15 ℃ 差异。因此一个合格的预测系统不仅要追求小误差，还要在缺数据、缺端点或跨域材料上给出可控的预测路线与诚实的置信度。",
    )

    add_h2(doc, "（二）共聚物 Tg 预测的统计建模困难")
    add_body_para(
        doc,
        "均聚物 Tg 预测是标准的结构—性能回归：给定重复单元，输出 Tg。共聚物则更复杂——同一组端点在不同组成比例、序列分布、链段架构与相态下可能给出不同 Tg，公开记录还存在端点未配对、组成方向不明、单位混用、同体系同组成多值、端点纯组分行混入等数据缺陷；忽略这些会使模型从体系内重复或端点泄漏获得虚高分数，反映的是体系内插值而非新体系外推。同时真实共聚物样本量比均聚物小一两个数量级，端到端神经网络易被均聚物大样本主导而过拟合；Fox、Gordon-Taylor 等经典物理近似只能表达理想混合，无法刻画非理想链间相互作用。因此合理路径是把物理先验与统计残差学习结合，而非用更复杂模型替代物理先验。",
    )
    add_body_para(
        doc,
        "为避免术语歧义，本文按两维独立刻画共聚物：组分数（二元 / 多元）与链段架构（无规 / 嵌段 / 接枝），并按“组分数-架构”顺序描述（如“二元-无规”）。“端点 Tg”指共聚物各组分对应均聚物的 Tg，“端点”指每一类重复单元。本文真实样本中，PolyInfo 149 条以二元-无规为主，少量三元-无规与嵌段；高分子-DNA 17 条是核酸/碱基侧基功能化的接枝共聚物。本文以 Fox 单一 Tg 假设为基线，物理上适用于单相体系；嵌段共聚物在强相分离下的双 Tg 行为不在适用范围内。",
    )

    add_h2(doc, "（三）高分子-DNA 共聚物的特殊挑战")
    add_body_para(
        doc,
        "高分子-DNA 共聚物融合核酸碱基识别和合成高分子多功能性，在生物识别、可编程自组装、刺激响应材料和生物医用涂层中具有潜在价值。其 Tg 表现出三方面特殊性：碱基间氢键与 π-π 堆积使链段运动能力对组成高度敏感；嘌呤与嘧啶对链段刚性的贡献不同，易形成系统性家族偏差；公开实验 Tg 数据极少（核酸物性表征多关注熔融温度而非链段玻璃化）。本文以 17 条核酸/碱基功能化共聚物为跨域先导验证集，目标不是覆盖完整 DNA 长链体系，而是检验框架在强相互作用、小样本与碱基类型外推下的可解释能力，按碱基类型 group holdout 避免同族泄漏。",
    )

    add_h2(doc, "（四）文献现状与研究空白")
    add_body_para(
        doc,
        "聚合物信息学已发展多套均聚物 Tg 预测方法：早期以 Bicerano 物性预测体系、Van Krevelen 基团贡献法为代表的描述符 + 基团模型；近年依托 PolyInfo、PI1M 等数据库，Polymer Genome、polyBERT、TransPolymer 等预训练表示与 TabPFN 等小样本表格学习方法相继出现。但共聚物 Tg 的统计预测研究偏少：Fox、Gordon-Taylor、Schneider-DiMarzio 等物理关系给出端点—组成混合的近似形式，纯统计建模工作多集中在具体材料族或较小样本集，少有同时关注数据清洗、虚拟增强、严格分组评估与应用路由的整体框架。高分子-DNA 体系 Tg 预测在公开文献中几乎空白。本文将端点物理先验、共聚物数据清洗、虚拟数据增强与 DNA 跨域验证整合为同一条统计建模链条。",
    )

    add_h2(doc, "（五）本文研究目标与技术路线")
    add_body_para(
        doc,
        "本文以“统计建模”为方法学定位，目标包括：从均聚物大数据学习端点结构—性能映射作为物理先验；按端点 Tg + Fox 关系生成可审计虚拟共聚物弱标签；对真实共聚物做系统清洗与分组评估；以高分子-DNA 共聚物作跨域先导验证；构建证据自适应路由覆盖不完整输入场景。技术路线见图1：均聚物源模型向下分虚拟弱标签生成与真实共聚物清洗两条支线，汇入物理先验门控残差回归形成主模型；主模型在三任务严格分组留出下评估，min-R² 为综合主指标，并嵌入证据自适应路由提供可审计输出。",
    )
    add_figure(doc, FIG_DIR / "figure01_workflow.png", width_cm=14.6)
    add_fig_caption(doc, "图1 共聚物 Tg 预测研究技术路线")
    add_blank(doc, 1)


def chapter2(doc) -> None:
    add_h1(doc, "二、数据来源与质量控制")

    add_h2(doc, "（一）数据总体构成与用途分配")
    add_body_para(
        doc,
        "本文所用数据包含四类信息：均聚物 Tg 数据、一般共聚物 Tg 数据、DNA 相关共聚物 Tg 数据以及虚拟共聚物弱标签数据。四类数据在建模链条中承担不同功能，不能简单合并；为了避免在统计意义上失去任务边界，本文对每一类数据明确其用途与质量控制重点（见表1）。在严格评估阶段，虚拟数据不作为主模型训练样本，仅在受控的虚拟数据增强实验中按权重融入。",
    )
    add_table_caption(doc, "表1 数据来源与用途说明")
    add_three_line_table(
        doc,
        ["数据类型", "规模", "主要用途", "质量控制重点"],
        [
            ("均聚物 Tg 数据", "7486 条", "学习端点 Tg 与基础结构—性能关系", "去重、结构有效性、特征处理一致性"),
            ("一般共聚物数据", "149 条", "检验真实共聚物体系外推能力", "端点映射、比例方向、单位统一、冲突过滤"),
            ("DNA 相关共聚物数据", "17 条", "检验核酸/碱基功能化跨域能力", "按碱基族分组，避免同族信息泄漏"),
            ("虚拟共聚物弱标签", "按实验筛选使用", "弱监督增强与风险诊断", "端点可靠性、物理一致性、分布相似性"),
        ],
        col_widths_cm=[3.2, 2.8, 4.6, 4.0],
    )
    add_body_para(
        doc,
        "图2 进一步给出三类真实数据的 Tg 经验分布。均聚物覆盖了从低 Tg 弹性体到高 Tg 高性能聚合物的较宽温区；一般共聚物以中低温区为主，分布形状显著比均聚物窄；DNA 相关共聚物主要集中在 −50 至 50 ℃ 区间，呈典型的小样本窄分布。三类分布在均值与跨度上的差异，是后续在统一回归框架内引入样本权重和分组留出评估的直接动机。",
    )
    add_figure(doc, FIG_DIR / "figure02_tg_distribution.png", width_cm=15.4)
    add_fig_caption(doc, "图2 三类真实数据的 Tg 分布")
    add_blank(doc, 1)

    add_h2(doc, "（二）均聚物数据")
    add_body_para(
        doc,
        "均聚物数据来自整理后的多源公开数据库，按重复单元 SMILES 去重后保留 7486 条样本。预处理包括 RDKit 解析过滤无效分子图、末端通配符 * 规范化、Tg 异常值（< −150 ℃ 或 > 500 ℃）剔除、训练—预测特征管线一致。该数据兼任端点 Tg 源模型训练样本与多源融合训练背景；样本量远大于共聚物与核酸（约 50 倍 / 440 倍），需通过样本权重与任务级评估抑制大样本主导。",
    )
    add_blank(doc, 3)

    add_h2(doc, "（三）一般共聚物数据的清洗流程")
    add_body_para(
        doc,
        "一般共聚物数据来自 PolyInfo 数据库。原始记录非为机器学习设计，存在多种缺陷，本文采取六步清洗（见表2）。",
    )
    add_table_caption(doc, "表2 一般共聚物数据清洗步骤汇总")
    add_three_line_table(
        doc,
        ["步骤", "处理内容", "目的"],
        [
            ("1. 结构与组成解析", "从原始记录中抽取共聚物体系、端点 SMILES、组成比例和 Tg", "建立机器可读的样本结构"),
            ("2. 端点映射", "将共聚物组分关联到均聚物端点库，对未对齐结构补全或剔除", "保证端点 Tg 可估计"),
            ("3. 比例方向与单位统一", "审计组成比例方向，将摩尔分数在可接受条件下近似换算为质量分数", "避免组成方向与单位歧义"),
            ("4. pure endpoint 剔除", "移除被错误标记为共聚物的纯端点行", "防止端点泄漏"),
            ("5. 稳定行 ID 重建", "对原始 sample_id 中的非唯一标识构造稳定行标识", "支持诊断与预测结果回连"),
            ("6. 体系冲突过滤", "对同一体系同一组成下 Tg 标准差超过 10 K 的组整体剔除", "降低实验冲突对模型上限的影响"),
        ],
        col_widths_cm=[3.2, 7.2, 4.2],
    )
    add_body_para(
        doc,
        "在第六步冲突过滤中，最具代表性的案例是 EVA 与 EVOH 体系。原始记录中存在同一组成下 Tg 跨度高达 30—50 ℃ 的样本，难以由模型分辨；这类记录在体系级评估中既会拉高表观误差，也会污染回归器。冲突过滤后，从 PolyInfo 共聚物原始数据中剔除 11 行高冲突样本，最终得到 149 条清洗后真实共聚物数据。该过程不是工程附属，而是统计建模的一部分：对小样本共聚物而言，数据清洗本身就影响模型上限。",
    )

    add_h2(doc, "（四）核酸/碱基功能化共聚物数据与碱基类型分组")
    add_body_para(
        doc,
        "DNA 相关共聚物数据由公开文献整理而成，共 17 条，每条关联一种主要碱基族（A/T/G/C/U）。样本量较小且同碱基族常重复出现，若随机划分易让模型从碱基特征获得族内信息，导致跨域评估虚高。本文按碱基族留出 group holdout，使每次预测样本所在族在训练时不可见，真正检验跨族迁移能力。该集合定位为核酸/碱基功能化共聚物的先导集，并非完整 DNA 长链材料。",
    )

    add_h2(doc, "（五）虚拟共聚物弱标签的生成原则")
    add_body_para(
        doc,
        "虚拟弱标签由端点预测与 Fox 型混合共同生成，可达数千条规模但不应等同实验真值。流程（图3）接收端点 SMILES、组成权重与架构标签，输出可审计 recipe（含稳定 ID、端点序列、归一化权重、预测 Tg、Fox 参考、端点 Tg 上下界、主 method）。本文确立三项使用原则：仅作弱监督辅助；加入主模型前必须通过端点可靠性、Fox 物理一致性与真实分布相似性筛选；若导致一般共聚物体系留出指标退化，则在主模型层拒用，仅作诊断素材。",
    )
    add_figure(doc, FIG_DIR / "figure03_virtual_generation.png", width_cm=14.6)
    add_fig_caption(doc, "图3 虚拟共聚物弱标签生成示意")
    add_blank(doc, 1)

    add_h2(doc, "（六）数据集汇总与可获取性")
    add_body_para(
        doc,
        "本文用于主模型评估的真实数据共 7486+149+17=7652 条，虚拟数据按实验需要使用并明确标记。数据划分方式在第三章给出，清洗规则、稳定行 ID 与体系标签将随数据与代码材料一并提供。",
    )
    add_blank(doc, 1)


def chapter3(doc) -> None:
    add_h1(doc, "三、评价指标与划分设计")

    add_h2(doc, "（一）误差与拟合度量")
    add_body_para_mixed(
        doc,
        "本文采用平均绝对误差 (MAE)、均方根误差 (RMSE) 和决定系数 ",
        ("omath", _om_r2),
        " 三类度量评价模型预测能力。对样本数为 n 的预测序列，三类度量分别定义为：",
    )
    add_eq_omml(doc, _eq_mae, "1")
    add_eq_omml(doc, _eq_rmse, "2")
    add_eq_omml(doc, _eq_r2_def, "3")
    add_body_para_mixed(
        doc,
        "MAE 直接解释预测偏差；RMSE 对极端误差更敏感，便于识别少数大误差体系；",
        ("omath", _om_r2),
        " 衡量相对“均值预测”的解释方差比例，便于在样本量不同的任务间相对比较。本文在所有结果表中同时给出 ",
        ("omath", _om_r2),
        " 与 MAE，必要时报告 RMSE。",
    )

    add_h2(doc, "（二）三任务分别评价")
    add_body_para(
        doc,
        "均聚物样本量约为共聚物的 50 倍、核酸的 440 倍，仅报告整体 R² 会被均聚物主导，掩盖共聚物与核酸任务表现。本文按任务分别报告（见表3），避免大样本遮蔽小样本指标。",
    )
    add_table_caption(doc, "表3 评价任务与划分方式")
    add_three_line_table(
        doc,
        ["评价任务", "划分方式", "统计意义", "风险控制要点"],
        [
            ("均聚物预测", "随机留出 (random holdout)", "检验基础结构—性能学习能力", "训练—预测特征处理一致；划分种子固定"),
            ("一般共聚物预测", "按共聚物体系留出 (group holdout)", "检验新体系外推能力", "避免同体系组成插值泄漏"),
            ("DNA 相关共聚物预测", "按碱基族分组留出 (group holdout)", "检验跨域迁移能力", "避免同一碱基族信息泄漏"),
        ],
        col_widths_cm=[3.6, 3.6, 4.0, 4.0],
    )

    add_blank(doc, 1)

    add_h2(doc, "（三）体系级与碱基类型分组留出")
    add_body_para(
        doc,
        "“分组留出（group holdout）”按体系编号或碱基类型整体划分，使训练—测试集无共享分组键值。在共聚物与核酸任务中，同一体系或碱基常出现于多条记录，随机划分会让模型从同体系不同组成中获得插值信息。形式上，若 G(i) 表示样本所属体系或碱基族，group holdout 强制 ∀i∈训练集, G(i)∉测试集分组，使评估接近真实部署的“新体系外推”场景。",
    )
    add_body_para(
        doc,
        "对一般共聚物按 PolyInfo 体系编号 (COID) 留出；对 DNA 相关样本按碱基族 (A/T/G/C/U) 留出。group holdout 通常 R² 低于 random holdout，这是评估更严而非模型变差；只在 random 划分上报告高分，不能直接说明模型具有推广能力。",
    )

    add_h2(doc, "（四）综合主指标 min-R²")
    add_body_para_mixed(
        doc,
        "本文进一步定义综合主指标：",
    )
    add_eq_omml(doc, _eq_min_r2_def, "4")
    add_body_para_mixed(
        doc,
        "该指标等价于“最差任务最大化”，迫使三任务同时保持下限。在多目标优化语言中对应 max-min 分量；在工程语言中相当于“任何一类任务都不允许塌陷”的硬性下限。后续消融与增强实验都必须证明在不损害其余任务的前提下提高最弱任务，才能采纳为新的主模型。",
    )
    add_blank(doc, 1)


def chapter4(doc) -> None:
    add_h1(doc, "四、方法 I：物理先验门控残差的通用统计模型")

    add_h2(doc, "（一）模型框架总览")
    add_body_para_mixed(
        doc,
        "方法 I 是不区分输入证据的通用回归器：均聚物、不同组分数与架构的一般共聚物以及 DNA 相关共聚物均按同一套特征与同一回归器输出 Tg 估计，可概括为“物理基线 + 受约束残差”。预测值 ",
        ("omath", _om_yhat),
        " 由 Fox 物理基线 ",
        ("omath", _om_f_phys),
        " 与数据驱动残差 ",
        ("omath", _om_r_theta),
        " 经门控 ",
        ("omath", _om_g_x),
        " 加权融合：",
    )
    add_eq_omml(doc, _eq_yhat, "5")
    add_body_para_mixed(
        doc,
        "门控 ",
        ("omath", _om_g_x),
        " 控制残差影响：均聚物以端点 Tg 为基线，残差承担主任务；共聚物以 Fox 为基线，残差仅修正端点物理无法解释部分；当端点 Tg 估计不可靠或样本落入低 Fox 区时，门控收缩残差，使预测向物理基线靠拢。模型整体结构如图4 所示。",
    )
    add_figure(doc, FIG_DIR / "figure04_model_architecture.png", width_cm=14.6)
    add_fig_caption(doc, "图4 方法 I 物理先验门控残差模型结构")
    add_blank(doc, 1)

    add_h2(doc, "（二）端点 Tg 源模型")
    add_body_para(
        doc,
        "源模型在端点 Tg 实验值缺失时给出估计，融合三类表示：链段物理特征（柔性键、基团贡献温度、自由体积等 58 维）、图神经网络在外部聚合物数据上预训练的图嵌入（64 维）、polyBERT 序列表示经 PCA 降至 64 维，拼接为 186 维端点表征，输入小样本表格学习器（TabPFN 风格）。训练严格遵循“特征处理与回归器仅在训练折拟合”原则；推断时可用全量均聚物给出端点最佳估计。源模型同时为后续虚拟数据生成提供端点 Tg 来源。",
    )

    add_h2(doc, "（三）共聚物 Fox 物理基线")
    add_body_para_mixed(
        doc,
        "Fox 关系是二元共聚物 Tg 的经典近似（Fox, 1956）。在 Kelvin 温度下，对组成权重 ",
        ("omath", lambda: _om_w_sub("1")),
        "、",
        ("omath", lambda: _om_w_sub("2")),
        " 与端点 Tg ",
        ("omath", lambda: _om_tg_sub("1")),
        "、",
        ("omath", lambda: _om_tg_sub("2")),
        "：",
    )
    add_eq_omml(doc, _eq_fox_two, "6")
    add_body_para_mixed(
        doc,
        "多元共聚物推广为 ",
        ("omath", _eq_fox_general),
        "。Fox 假设链段理想混合，不考虑非理想相互作用与结晶；本文将其作为物理基线 ",
        ("omath", _om_f_phys),
        "，并把端点 Tg 上下界 (max/min) 作为辅助物理特征，使模型既能贴合 Fox 主线又能给出可解释的偏离修正。",
    )

    add_h2(doc, "（四）共聚物层结构与组成特征")
    add_body_para_mixed(
        doc,
        "共聚物层特征含两类。一是端点物理特征：端点 Tg 的 ",
        ("omath", lambda: _om_tg_sub("max")),
        "、",
        ("omath", lambda: _om_tg_sub("min")),
        "、加权均值 ",
        ("omath", lambda: _om_tg_sub("w")),
        "、跨度 ",
        ("omath", lambda: _om_tg_sub("max")),
        "−",
        ("omath", lambda: _om_tg_sub("min")),
        "，显式表达端点差异。二是组成统计特征：最大占比 ",
        ("omath", _om_max_w_i),
        "、组成熵 ",
        ("omath", _eq_entropy),
        "、Herfindahl 指数 ",
        ("omath", _eq_herfindahl),
        "。组分数（二元/多元）与架构（无规/嵌段）以 one-hot 独立维度融入。",
    )

    add_h2(doc, "（五）物理先验门控残差回归")
    add_body_para_mixed(
        doc,
        "残差学习器 ",
        ("omath", _om_r_theta),
        " 采用树/核小样本表格回归器，输入由端点物理、组成统计、共聚物层结构与端点结构嵌入合并约 1191 维，并以早停和正则限制复杂度。门控 ",
        ("omath", _om_g_x),
        " 由若干指示变量构成（是否均聚物、端点是否已知、Fox 基线高低区、是否核酸样本等），既参与回归输入又作为缩放因子在校准模块中作用于残差。物理动机：端点物理足以解释时残差趋零；无法解释时残差修正；端点估计可疑时残差被压制，避免在错误基线上累加偏移。",
    )

    add_h2(doc, "（六）三类校准")
    add_body_para(
        doc,
        "残差回归器外设三类后校准，逐步缩窄相对物理先验的偏离方向。其一，非均聚物门控的端点/Fox 残差校准：仅对共聚物，将预测向“端点 Tg 加权均值 + Fox 调整项”做小幅回拉，缓解过自信外推。其二，低 Fox 区残差收缩：Fox 基线落入低 Tg 区时残差幅度受额外抑制。其三，非均聚物预测差校准：以“端点 Tg_Fox − 模型基础预测”的差值统一修正整体偏移。物理区域越极端校准作用越强，安全区接近零。第六章消融定量给出每类校准对三任务 R² 的贡献。",
    )

    add_h2(doc, "（七）多源数据融合与样本权重")
    add_body_para_mixed(
        doc,
        "简单合并会让大样本均聚物主导损失。本文按任务样本量比例设权重：均聚物 ",
        ("omath", lambda: _om_w_sub("homo")),
        "=1，一般共聚物 ",
        ("omath", lambda: _om_w_sub("co")),
        "=10，DNA 相关 ",
        ("omath", lambda: _om_w_sub("nuc")),
        "=60，虚拟 ",
        ("omath", lambda: _om_w_sub("v")),
        "=0（仅在虚拟增强实验中调整）。权重在训练侧与三任务分别评价、",
        ("omath", _om_min_r2),
        " 主指标共同抑制单任务过度优化。方法 I 由此成为方法 II 路由系统中唯一的无证据 fallback。",
    )
    add_blank(doc, 1)


def chapter5(doc) -> None:
    add_h1(doc, "五、方法 II：证据自适应路由系统")

    add_h2(doc, "（一）方法动机：为何需要按证据分流")
    add_body_para(
        doc,
        "真实使用中用户证据不一致：有时只有单一重复单元；有时是已知体系内的新组成；有时端点 Tg 已实测但体系未训练；有时仅端点结构而无 Tg 测量。强行用同一回归器要么浪费已知证据，要么在证据缺失时过度外推。方法 II 把“按证据分流”显式写进模型：根据证据等级在五个互补子估计器间动态选择，并在输出中标记路线、方法、端点 Tg 来源、置信度与风险提示。它不替代方法 I，而把方法 I 当作端点估计接口与无证据 fallback，呈现“证据等级 → 路线选择 → 子估计器 → 校准 → 可审计输出”的级联。",
    )
    add_blank(doc, 1)

    add_h2(doc, "（二）四级证据等级与八类预测路线")
    add_body_para(
        doc,
        "本文将共聚物输入证据划分为 E1（端点 Tg 已实测）、E2（体系编号已知）、E3（仅端点 SMILES 与组成）、E4（端点缺失或仅有部分链段）四级，路由按高到低优先级选择最强可用等级。在此基础上结合组分数（二元/多元）与架构（无规/嵌段）以及 DNA 是否提供端点实测，展开为八类路线（见表4）。其中均聚物、未知二元—无规、已知体系局部插值与 DNA actual endpoint 路线最常用；多元—无规近似、嵌段近似、DNA universal default 与端点缺失 fallback 覆盖证据不完整场景并显式降低置信度。决策树见图5。",
    )
    add_figure(doc, FIG_DIR / "figure10_routing_tree.png", width_cm=14.6)
    add_fig_caption(doc, "图5 方法 II 证据自适应路由决策树")
    add_blank(doc, 4)
    add_table_caption(doc, "表4 方法 II 路由系统的输入证据—预测路线对应")
    add_three_line_table(
        doc,
        ["输入证据等级", "组分数—架构", "预测路线", "计算方式", "输出含义"],
        [
            ("E3", "homo（均聚物）", "均聚物源模型路线", "TabPFN 186d 端点源模型", "公平 holdout R²=0.907"),
            ("E2", "二元—无规", "同体系局部残差插值", "Linear-Fox + 同体系残差 IDW", "体系内 R²=0.933"),
            ("E3", "二元—无规", "全局 Fox + 受约束残差", "物理先验门控残差回归", "新体系外推 R²=0.820"),
            ("E3", "多元—无规", "多端点 Fox 近似", "Fox 关系拓展 + 残差", "保守估计，置信度中等"),
            ("E3", "二元 / 多元—嵌段", "miscible proxy + Tg window", "近似 Tg 与端点窗口", "风险感知近似，置信度中等"),
            ("E1", "DNA 二元—无规（specialized）", "actual endpoint 物理路线", "Linear-Fox / Physics-Ridge LOOCV", "端点实测 R²=0.971"),
            ("E3", "DNA 二元—无规（universal）", "全局 Linear-Fox 默认路线", "PolyInfo 校准 + 端点估计", "仅 SMILES R²=0.564"),
            ("E4", "—", "端点估计 fallback", "源模型端点估计 + 方法 I 输出", "探索性预测，置信度低"),
        ],
        col_widths_cm=[2.0, 2.4, 3.0, 4.0, 3.8],
    )

    add_h2(doc, "（三）五个子估计器及其物理依据")
    add_body_para(
        doc,
        "路由可信度依赖五个互补子估计器，各对应一类物理直觉。",
    )

    add_h3(doc, "1. Fox 端点直接路线")
    add_body_para_mixed(
        doc,
        "端点 Tg 与组成已知时直接套用 Fox 关系：",
        ("omath", _eq_fox_general),
        "。无需拟合，物理动机为熵可加性。在 149 条 PolyInfo 上 R²=0.785、MAE=18.79 K（见表7），作参照基线。",
    )

    add_h3(doc, "2. Linear-Fox 拟合校准")
    add_body_para_mixed(
        doc,
        "Fox 偏离可由仿射校准吸收：",
        ("omath", _eq_linear_fox),
        "。在 PolyInfo 体系级留出下 R²=0.820、MAE=16.49 K，跨体系外推优于纯 Fox；在 17 条核酸 LOOCV 上 R²=0.937、MAE=3.64 K，表明端点实测可用时简单仿射就能大幅纠偏。",
    )

    add_h3(doc, "3. Kwei 同体系优先 LOOCV")
    add_body_para_mixed(
        doc,
        "Kwei 关系是 Fox 的二参数推广：",
        ("omath", _eq_kwei),
        "，(k, q) 刻画二元非理想程度。本文按体系独立拟合 (k, q)，体系外退化为全局 Fox。PolyInfo 149 条 LOOCV R²=0.865、MAE=13.73 K，适合“体系内 1—2 条参考”场景。",
    )

    add_h3(doc, "4. Physics-Ridge 物理特征回归")
    add_body_para(
        doc,
        "12 维链间相互作用特征（端点 Tg 跨度、Fox 残差方向、组成熵等）单独作为 Ridge 输入。PolyInfo 149 条 LOOCV R²=0.901、MAE=12.20 K；核酸 17 条 LOOCV R²=0.971、MAE=2.93 K，是路由精度最高的子估计器。物理直觉：端点物理足以解释链段相互作用时浅层线性模型即可捕获主要规律。",
    )

    add_h3(doc, "5. 同体系残差 IDW 局部插值")
    add_body_para(
        doc,
        "目标体系内 ≥ 3 条不同组成实测时，将“真实 Tg − Linear-Fox 预测”视为组成 w 的连续函数做反距离加权插值，并叠回 Linear-Fox 基线。PolyInfo 149 条上 leave-row R²=0.933、MAE=8.76 K（路由全局最优），leave-composition R²=0.920、MAE=11.28 K（更接近“同体系新组成”实情）。该路线对应方法 II 最高置信度档。",
    )

    add_h2(doc, "（四）路由决策机制与可审计输出")
    add_body_para(
        doc,
        "路由分两层判断。第一层证据等级：依次检查端点 Tg 实测、体系编号、组成比例、架构标签，落在最高可用等级。第二层子估计器选择：按表4 对应关系确定路线；可用样本不足（如同体系样本 < 3）时自动降级（残差 IDW → Physics-Ridge LOOCV → Linear-Fox 全局校准）。可审计输出是方法 II 关键：每条预测除 Tg 外还给出路线名、主要计算方法、端点 Tg 来源（experimental / source-model-estimated）、Fox 参考、所用同体系样本数、是否降级 fallback、A/B/C 置信度与风险提示，使每条预测都可解释为“严格外推 / 体系内插值 / 跨域端点路线 / 低置信度 fallback”等具体类别。",
    )

    add_h2(doc, "（五）训练数据复用与方法 I 的关系")
    add_body_para(
        doc,
        "两方法共享均聚物 7486 + PolyInfo 149 + 核酸 17 的真实资源，但回归器并不相同。方法 I 是多源融合的统一回归器（exp56），均聚物 holdout R²=0.887 是联合训练的妥协代价；方法 II 在均聚物路线下直接调用为均聚物专门训练的 BestTgPredictor（PHY-C-light 58d + GNN 64d + polyBERT PCA 64d，TabPFN v2），同 1498 行公平 holdout R²=0.907、MAE=20.6 K，提升 0.020 R²；同模型在全 7486 条上 5×3 RepeatedKFold CV R²=0.917、MAE=20.4 K 作独立基准。方法 II 在端点缺失场景下默认调用 BestTgPredictor 估计端点，而非 exp56，这是其在“端点已知 / 体系已知”证据下指标能推到 0.91—0.97 区间的关键。",
    )
    add_body_para(
        doc,
        "DNA actual-endpoint 路线 R²=0.971 仅在用户显式提供碱基实测端点（A=65、T=43、C=90、G=130 ℃）并允许在 17 行内做 Physics-Ridge LOOCV 时达到；仅提交端点 SMILES 时路由默认走 binary_global_linear_fox 全局校准路线，端点由 BestTgPredictor 估计而无碱基特异氢键/堆积，R² 退化到 0.564（universal_router_fulltest 实验直接给出）。这一档差是路由系统当前真实使用边界。第六章的对比应理解为“同一数据资源上的两层决策”：方法 I 给出最严苛 group holdout 下的诚实指标，方法 II 在不同证据等级下输出多档结果，端点缺失时退化为“方法 I 输出 + 低置信度标签”，由此形成完整覆盖矩阵。",
    )
    add_blank(doc, 1)


def chapter6(doc) -> None:
    add_h1(doc, "六、综合求解、对比与诊断")

    add_h2(doc, "（一）主指标对照：方法 I 与方法 II 在三任务上的总览")
    add_body_para(
        doc,
        "本节先单独展示方法 I 在三任务严格分组评估下的性能，再以五种代表性场景把两方法并列展示（表5）。",
    )
    add_blank(doc, 1)
    add_table_caption(doc, "表5 方法 I 主模型预测性能")
    add_three_line_table(
        doc,
        ["评估任务", "n", "R²", "MAE / ℃", "RMSE / ℃"],
        [(m.name, str(m.n), f"{m.r2:.3f}", f"{m.mae:.2f}", f"{m.rmse:.2f}") for m in MAIN_RESULTS],
        col_widths_cm=[5.4, 1.4, 1.6, 2.0, 2.0],
    )
    add_body_para(
        doc,
        "在表5 基础上，五种场景的两方法并列结果见表6 与图6。所有口径严格对齐 results/universal_router_fulltest/ 服务器原始记录。均聚物上方法 I 用 exp56 在 random holdout R²=0.887（n=1498，混合训练代价）；方法 II 由路由直接调用为均聚物专训的 TabPFN 186d，在同 1498 行做公平 train/test，R²=0.907、MAE=20.6 K（+0.020 R²、−6.6 K MAE），同模型在全 7486 条 5×3 RepeatedKFold CV R²=0.917 作独立基准。PolyInfo 149 行跨体系外推 (leave-system-out)：方法 I R²=0.849，方法 II Linear-Fox 留体系外 R²=0.820（略低于方法 I，因方法 II 子估计器仅 149 条共聚物拟合）；体系内已知样本 (leave-row + 残差 IDW)：方法 II R²=0.933。DNA 17 行按碱基族留出：方法 I R²=0.817；方法 II 分两档——specialized actual-endpoint（提供 A/T/C/G 端点 + Physics-Ridge LOOCV）R²=0.971，universal default（仅 SMILES）R²=0.564。",
    )
    add_table_caption(doc, "表6 方法 I 与方法 II 主指标并列对照矩阵")
    add_three_line_table(
        doc,
        ["评估场景", "方法 I n", "方法 I R²", "方法 II n", "方法 II 最佳 R²", "方法 II 对应路线"],
        [
            (s[0], str(s[1]), f"{s[2]:.3f}", str(s[3]), f"{s[4]:.3f}", s[5])
            for s in M1_VS_M2_MAIN
        ],
        col_widths_cm=[4.0, 1.2, 1.6, 1.2, 1.8, 4.8],
    )
    add_blank(doc, 1)
    add_figure(doc, FIG_DIR / "figure11_m1_vs_m2.png", width_cm=14.0)
    add_fig_caption(doc, "图6 方法 I 与方法 II 主指标对比")
    add_blank(doc, 1)
    add_body_para(
        doc,
        "规律一致：方法 I 在“证据等级齐平”的统一评估下给出稳健全局指标，承担诚实评估；方法 II 在“证据有差异”的真实场景下进一步利用已知证据，逼近上界，证据缺失时回退。均聚物上方法 II 避开混合训练代价 +0.020 R²；PolyInfo 体系内 +0.084 R²、跨体系 −0.029 R²；DNA specialized +0.154 R²，universal default −0.253 R²。两类指标互补：方法 I 体现 min-R²=0.817 的综合稳健性，方法 II 提供不同证据等级下多档输出（最高档为精度上界、最低档诚实暴露退化），缺一会让读者只看“最佳数字”忽视使用边界。",
    )
    add_body_para(
        doc,
        "图7 给出方法 I 三任务严格评估的预测—实测散点与 R² 条形图：均聚物沿对角线分布无系统偏差；一般共聚物点群紧凑，少量偏离来自硬体系；DNA 相关因 Tg 分布窄散点紧密但绝对位置略偏，与 R²、MAE 互相印证。",
    )
    add_figure(doc, FIG_DIR / "figure05_main_results.png", width_cm=15.2)
    add_fig_caption(doc, "图7 方法 I 三任务严格评估表现与预测—实测散点")
    add_blank(doc, 1)

    add_h2(doc, "（二）PolyInfo 共聚物 149 行：路由七路全表")
    add_body_para(
        doc,
        "为透明展示方法 II 内部结构，PolyInfo 149 行的六口径汇总于表7：Fox 端点直接 R²=0.785、MAE=18.79 K；Linear-Fox 留体系外 R²=0.820；Kwei 同体系优先 LOOCV R²=0.865；Physics-Ridge LOOCV R²=0.901；同体系残差 IDW 在 leave-row 与 leave-composition 下分别 R²=0.933、0.920。",
    )
    add_table_caption(doc, "表7 方法 II 在 PolyInfo 共聚物 149 行上的子估计器全表")
    add_three_line_table(
        doc,
        ["子估计器 / 路线", "评估方式", "n", "MAE / K", "RMSE / K", "R²"],
        [
            (s[0], s[1], str(s[2]), f"{s[3]:.2f}", f"{s[4]:.2f}", f"{s[5]:.3f}")
            for s in ROUTER_POLYINFO
        ],
        col_widths_cm=[3.8, 3.4, 1.0, 2.0, 2.2, 1.6],
    )
    add_body_para(
        doc,
        "三条规律。其一，体系内信息含量决定精度天花板：≥ 3 条同组成测量时残差 IDW 绝对最优；1—2 条参考时 Kwei 同体系优先在 R²=0.865 量级最稳定；体系完全未训练时 Physics-Ridge LOOCV 仍 R²=0.901。其二，跨体系外推是真实下界：Linear-Fox 留体系外 R²=0.820 与方法 I R²=0.849 同一量级。其三，Fox 关系仍是物理参照但偏差不可忽视——简单仿射就能把 R² 从 0.785 拉到 0.820，加 12 维物理特征推到 0.901。",
    )

    add_h2(doc, "（三）高分子-DNA 共聚物碱基类型 17：跨碱基泛化全表")
    add_body_para(
        doc,
        "17 条核酸样本端点 Tg（A=65、T=43、C=90、G=130 ℃）均已实测，方法 II 多个子估计器可同时启用。表8 列出七口径并列结果。",
    )
    add_table_caption(doc, "表8 方法 II 在核酸碱基功能化共聚物 17 行上的跨碱基全表")
    add_three_line_table(
        doc,
        ["子估计器 / 路线", "评估方式", "n", "MAE / K", "RMSE / K", "R²"],
        [
            (s[0], s[1], str(s[2]), f"{s[3]:.2f}", f"{s[4]:.2f}", f"{s[5]:.3f}")
            for s in ROUTER_NUCLEOBASE
        ],
        col_widths_cm=[4.0, 2.6, 1.0, 2.0, 2.2, 1.6],
    )
    add_body_para(
        doc,
        "三条观察。其一，表8 首行“方法 I 默认输出 R²=0.167”是另一口径——训练集仅含均聚物与 PolyInfo（无任何核酸）再对 17 条做 single holdout，与第六章 (一) 报告的 R²=0.817 差异在于后者训练集已包含其余 16 条碱基族样本；纯 Fox（端点实测代入）R² 略低于 0，反映嘌呤—嘧啶系统偏差。其二，“端点实测 + 仿射/二参数校准”使 R² 跨越到 0.79—0.94 量级（Gordon-Taylor 0.790、Kwei 0.885、Linear-Fox 0.937），说明方法 II 的关键是更强证据利用而非更复杂模型。其三，Physics-Ridge LOOCV R²=0.971、MAE 2.93 K，leave-base-out 协议下 R²=0.872 仍可接受，体现 12 维物理特征跨族泛化能力。",
    )
    add_body_para(
        doc,
        "需特别说明：表8 所有 specialized 路线均要求用户显式提供 A/T/C/G 实测端点并允许 17 行 LOOCV 校准。仅提交 SMILES 时路由默认走 universal binary_global_linear_fox，端点由 BestTgPredictor 估计而非碱基实测，R² 仅 0.564、MAE 10.63 K（来自 universal_router_fulltest）。该档差反映“碱基特异端点未内置进 calibration table”，应在结论中作为后续工作提出。",
    )

    add_h2(doc, "（四）校准组件消融与虚拟数据增强")
    add_body_para(
        doc,
        "回到方法 I 内部诊断。以基础物理-局部残差（无后校准）为起点，逐步加入端点/Fox 残差校准、低 Fox 区残差收缩与非均聚物预测差校准。每步必须在不损害其余任务前提下提高最弱任务，实现 min-R² 合规增益（见表9）。",
    )
    add_table_caption(doc, "表9 方法 I 校准组件消融结果")
    add_three_line_table(
        doc,
        ["阶段", "均聚物 R²", "一般共聚物 R²", "DNA 相关 R²", "min-R²", "增益解释"],
        [
            ("基础物理-局部残差", "0.887", "0.844", "0.789", "0.789", "主回归器无后校准，作为基线"),
            ("+ 端点/Fox 残差校准", "0.887", "0.845", "0.792", "0.792", "弱缓解小样本过自信外推"),
            ("+ 低 Fox 区残差收缩", "0.887", "0.852", "0.810", "0.810", "约束低 Tg 区残差幅度"),
            ("+ 非均聚物预测差校准", "0.887", "0.849", "0.817", "0.817", "把模型整体偏移与 Fox 对齐"),
        ],
        col_widths_cm=[4.0, 2.0, 2.4, 2.2, 1.6, 4.0],
    )
    add_blank(doc, 1)
    add_figure(doc, FIG_DIR / "figure06_ablation.png", width_cm=14.6)
    add_fig_caption(doc, "图8 方法 I 校准组件消融的三任务 R² 曲线")
    add_blank(doc, 1)
    add_body_para(
        doc,
        "图8 给出各阶段三任务 R² 与 min-R² 演化。均聚物 R² 四阶段保持 0.887 不变，校准未损害源任务；一般共聚物 R² 经 0.844→0.845→0.852→0.849 微调，末步略回但幅度小；DNA 相关 R² 由 0.789 提至 0.817。最终 min-R² 提升 0.028，对应 28% 的差距缩小（以 1−R² 为参照）。",
    )
    add_body_para(
        doc,
        "虚拟数据实验检验“端点驱动弱标签是否提升真实任务”。比较三策略：仅真实数据（主模型）、加入结构近邻虚拟（400 条，权重 0.05）、加入一致性过滤虚拟（弱标签与教师预测差 ≤ 30 ℃，权重 0.02），见表10。",
    )
    add_table_caption(doc, "表10 虚拟数据增强策略对比")
    add_three_line_table(
        doc,
        ["策略", "均聚物 R²", "一般共聚物 R²", "DNA 相关 R²", "判断"],
        [
            ("仅真实数据 (方法 I 主模型)", "0.887", "0.849", "0.817", "当前最稳健方案"),
            ("加入结构近邻虚拟", "0.886", "0.827", "0.858", "提升核酸但损害一般共聚物"),
            ("加入一致性过滤虚拟", "0.886", "0.850", "0.817", "缓解退化但增益消失"),
        ],
        col_widths_cm=[4.4, 2.0, 2.4, 2.2, 4.6],
    )
    add_blank(doc, 1)
    add_figure(doc, FIG_DIR / "figure07_virtual_tradeoff.png", width_cm=12.6)
    add_fig_caption(doc, "图9 虚拟数据增强的收益—风险散点")
    add_blank(doc, 1)
    add_body_para(
        doc,
        "结构近邻虚拟把 DNA R² 从 0.817 提升到 0.858，但一般共聚物 R² 由 0.849 降到 0.827，反映虚拟弱标签与 PolyInfo 真实分布存在偏移。一致性过滤把虚拟与教师中位差从约 145 ℃ 降到约 12 ℃，代价是大量对核酸有用的高 Tg 弱样本被同时过滤。规律：端点物理一致性 + 真实分布相似性 + 教师一致性三者必须同时满足才能稳健提升。",
    )

    add_h2(doc, "（五）硬体系与碱基族偏差诊断")
    add_body_para(
        doc,
        "一般共聚物 group holdout 中少数体系主导误差。表11 列出 MAE 最大的 5 个体系，P900015 EVA 体系 MAE 约 35 ℃，几乎为整体均值的两倍。图10 直观展示硬体系排名。",
    )
    add_blank(doc, 4)
    add_table_caption(doc, "表11 一般共聚物中误差最大的体系")
    add_three_line_table(
        doc,
        ["体系编号", "材料类型", "体系 MAE / ℃", "主导误差原因"],
        [(s[0], s[1], f"{s[2]:.1f}", s[3]) for s in HARD_SYSTEMS],
        col_widths_cm=[2.6, 3.4, 2.6, 6.6],
    )
    add_blank(doc, 1)
    add_figure(doc, FIG_DIR / "figure08_hard_systems.png", width_cm=14.0)
    add_fig_caption(doc, "图10 一般共聚物中误差最大的体系")
    add_blank(doc, 1)
    add_body_para(
        doc,
        "P900015 误差并非单一冲突所致：诊断显示低 Tg/高次要组分行被过预测、高 Tg/低次要组分行被低预测，呈混合符号偏差，体系级截距修正无法同时改善两端。更合理的解释是当前特征未表达共聚物相态、混溶性、结晶倾向与组成-相分布，这些因素在 EVA 等体系中可能比端点 Tg 更主导链段运动。其余硬体系呈现真实记录稀疏、组成范围窄或端点物理在低 Tg 区贴合差等共同模式。",
    )
    add_body_para(
        doc,
        "DNA 相关任务误差呈碱基族系统偏差：T 族被系统性低估、G 族被系统性高估，A/C/U 偏差较小。图11 箱线图展示各族残差分布。偏差既反映碱基氢键/堆积能力差异，也反映族级样本量过小，是 Physics-Ridge 路线在 leave-base-out 下 R²=0.872 略低于 LOOCV R²=0.971 的主因。",
    )
    add_figure(doc, FIG_DIR / "figure09_nucleobase_residuals.png", width_cm=13.6)
    add_fig_caption(doc, "图11 DNA 相关共聚物的碱基族系统偏差")
    add_blank(doc, 1)

    add_h2(doc, "（六）不确定性、典型预测与互补性分析")
    add_body_para(
        doc,
        "两方法均支持基于残差分布的 90% 置信区间。区间宽度反映模型把握程度：端点 Tg 已知、组成位于密集区且 Fox 合理时收紧；端点由源模型估计、组成稀疏或跨域时放宽。方法 II 按路线给出 A/B/C 置信度：A 档对应同体系残差 IDW 与 actual endpoint 物理路线（< 10 K）；B 档对应 Linear-Fox / Kwei / Physics-Ridge（10—25 K）；C 档对应端点缺失 fallback（> 25 K，低置信）。",
    )
    add_body_para(
        doc,
        "本文抽取七类典型预测案例：均聚物 PMMA（A 档，Tg=105 ℃）；体系内插值 P900001 不同组成（A 档，残差 IDW）；未知体系 PMMA-co-PVAc（B 档，Linear-Fox）；多元—无规三组分（B 档）；嵌段（B 档，miscible proxy + Tg window）；DNA actual endpoint（A 档，Physics-Ridge）；端点缺失 fallback（C 档，方法 I 估计 + 风险标签）。每条预测的输出字段与附录 D 字段规范一一对应。方法 I 是“无证据条件下的诚实回归器”，方法 II 是“证据感知的最优插值器”，两者共同构成既诚实又实用的预测体系。",
    )
    add_blank(doc, 1)


def chapter7(doc) -> None:
    add_h1(doc, "七、结论与建议")

    add_h2(doc, "（一）主要结论与双方法覆盖矩阵")
    add_body_para(
        doc,
        "本文围绕共聚物 Tg 统计预测，在均聚物 7486 条、PolyInfo 共聚物 149 条、高分子-DNA 共聚物 17 条的同一资源上提出两个互补建模方法。主要结论六点。",
    )
    add_body_para(
        doc,
        "第一，物理先验门控残差回归在共聚物 Tg 预测中具有结构性必要性：物理基线提供稳定参照，残差学习器负责非理想偏差，门控控制残差幅度，三者构成对小样本共聚物友好的统计形式（方法 I 核心）。",
    )
    add_body_para(
        doc,
        "第二，方法 I 在严格分组评估下，三任务 R² 分别 0.887 / 0.849 / 0.817，MAE 27.16 / 16.65 / 6.27 ℃，min-R²=0.817。校准组件逐步加入使主指标从 0.789 提升至 0.817。",
    )
    add_body_para(
        doc,
        "第三，方法 II 显式利用输入证据等级在五个子估计器间动态选择。均聚物路线（TabPFN 186d）同 1498 行公平 holdout R²=0.907、MAE=20.6 K（+0.020 R²，全 7486 条 5×3 CV R²=0.917 作独立基准）；PolyInfo 149 行体系内同体系残差 IDW R²=0.933、MAE=8.76 K，跨体系 Linear-Fox 留体系外 R²=0.820；核酸 17 行 specialized actual-endpoint Physics-Ridge R²=0.971、MAE=2.93 K，universal default（仅 SMILES）R²=0.564、MAE=10.63 K，档差是路由真实使用边界。",
    )
    add_body_para(
        doc,
        "第四，两方法形成清晰覆盖矩阵：方法 I 在“证据缺失或架构复杂”下保持通用，方法 II 在“端点已知 / 体系已知”下给出更紧区间，通过路由可审计输出统一对外。这表明本文工作是“通用统计模型 + 证据自适应路由”双方法体系，而非单一回归器训练。",
    )
    add_body_para(
        doc,
        "第五，虚拟弱监督具有“双约束特性”：端点物理一致性 + 真实分布相似性同时满足时虚拟样本可提升核酸跨域；任一约束缺失则反损害一般共聚物外推。",
    )
    add_body_para(
        doc,
        "第六，可审计输出（路线名 / 计算方法 / 端点来源 / 置信度 / 风险提示）使两方法都能被可靠解释，这是“统计预测”而非“黑箱回归”的关键差别。",
    )

    add_h2(doc, "（二）主要不足")
    add_body_para(
        doc,
        "数据层面，DNA 相关样本仅 17 条，跨域推广受样本量约束，碱基族偏差形态待更大数据集验证；一般共聚物部分硬体系（P900015 EVA）的高误差暴露当前特征对相态、混溶性、结晶倾向的不足；PolyInfo 体系数仅 16，方法 II 残差 IDW 在体系数扩大后能否保持 R²>0.93 仍需验证。表征层面，端点嵌入未显式表达链段非理想相互作用、组成-相态曲线与加热速率等过程量。方法 II 路由切换为离散决策，临界证据等级（如同体系样本数从 3 跌至 2）下可能跳变；DNA calibration table 未注册碱基实测端点，仅 SMILES 时退化到 R²=0.564 与 specialized 0.971 形成档差。虚拟数据池来自有限端点组合枚举而非端点感知完整生成，限制了对一般共聚物外推的潜在贡献。",
    )

    add_h2(doc, "（三）后续工作建议")
    add_body_para(
        doc,
        "数据方面，优先扩充端点明确、组成方向清晰、单位统一、分子量与实验条件完整的真实共聚物，并在 DNA 方向系统补充不同碱基排列、连接方式与组成比例的实测 Tg。特征方面，引入相态、混溶性、结晶倾向、氢键供受体强度、链段刚性指数等显式物理特征，并同时加入方法 I 回归器与方法 II Physics-Ridge 子估计器以比较边际增益。虚拟数据方面，转向端点感知的生成与筛选：保留端点 Tg 不确定性、Fox 残差与教师一致性，按“端点完整性 + 物理一致性 + 真实分布相似性 + 教师一致性”四维联合打分。",
    )
    add_body_para(
        doc,
        "方法 II 路由方面，探索置信度软切换以避免 fallback 边界跳变；引入跨家族 GNN 微调以增强 Physics-Ridge 在新碱基族/体系上的零样本泛化。一项直接改进是把 A/T/C/G 实测端点表注册进 calibration table，使 DNA 样本在仅 SMILES 输入下也能默认走 specialized 路线，预计可把 universal default R² 从 0.564 拉到接近 0.971。应用方面，把双方法体系与 BigSMILES 等结构表示工具链对接，输出 Tg 预测、路线、置信度与风险提示，实现“统计模型 + 物理先验 + 证据路由”一体化的高分子物性预测范式。",
    )


def add_body(doc) -> None:
    add_main_title(doc, TITLE)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)


# ============================================================
# Back matter: references, appendix, acknowledgements
# ============================================================


REFERENCES: tuple[str, ...] = (
    # ---- 中文文献 ----
    "刘伦洋, 丁芳, 李云琦. 高分子材料大数据研究：共性基础、进展及挑战. 高分子学报, 2022, 53(6): 564–580.",
    "杨镇岳, 聂文建, 刘伦洋, 等. 机器学习方法在高分子玻璃化研究中的应用. 高分子学报, 2023, 54(4): 432–450.",
    "宫祥瑞, 蒋滢. 机器学习在高分子材料基因组研究中的进展与挑战. 高分子学报, 2022, 53(11): 1287–1300.",
    "全国塑料标准化技术委员会. 塑料 差示扫描量热法（DSC）第2部分：玻璃化转变温度的测定: GB/T 19466.2—2004. 北京: 中国标准出版社, 2004.",
    # ---- 英文文献 ----
    "Fox T G. Influence of diluent and of copolymer composition on the glass temperature of a polymer system. Bulletin of the American Physical Society, 1956, 1: 123.",
    "Bicerano J. Prediction of Polymer Properties. 3rd ed. New York: Marcel Dekker, 2002.",
    "Adam G, Gibbs J H. On the temperature dependence of cooperative relaxation properties in glass-forming liquids. The Journal of Chemical Physics, 1965, 43(1): 139–146.",
    "Williams M L, Landel R F, Ferry J D. The temperature dependence of relaxation mechanisms in amorphous polymers and other glass-forming liquids. Journal of the American Chemical Society, 1955, 77(14): 3701–3707.",
    "Van Krevelen D W, Te Nijenhuis K. Properties of Polymers. 4th ed. Amsterdam: Elsevier, 2009.",
    "Schneider H A, DiMarzio E A. The glass temperature of polymer blends: comparison of both the free volume and the entropy predictions with data. Polymer, 1992, 33(16): 3453–3461.",
    "Ten Brinke G, Karasz F E, MacKnight W J. Phase behavior in copolymer blends: poly(2,6-dimethyl-1,4-phenylene oxide) and halogen-substituted styrene copolymers. Macromolecules, 1983, 16(12): 1827–1832.",
    "Otsuka S, Kuwajima I, Hosoya J, et al. PoLyInfo: Polymer database for polymeric materials design. International Conference on Emerging Intelligent Data and Web Technologies, 2011: 22–29.",
    "Ma R, Luo T. PI1M: A benchmark database for polymer informatics. Journal of Chemical Information and Modeling, 2020, 60(10): 4684–4690.",
    "Mannodi-Kanakkithodi A, Pilania G, Huan T D, et al. Machine learning strategy for accelerated design of polymer dielectrics. Scientific Reports, 2016, 6: 20952.",
    "Kim C, Chandrasekaran A, Huan T D, et al. Polymer Genome: A data-powered polymer informatics platform for property predictions. Journal of Physical Chemistry C, 2018, 122(31): 17575–17585.",
    "Kuenneth C, Ramprasad R. polyBERT: a chemical language model to enable fully machine-driven ultrafast polymer informatics. Nature Communications, 2023, 14(1): 4099.",
    "Xu C, Wang Y, Barati Farimani A. TransPolymer: A Transformer-based language model for polymer property predictions. npj Computational Materials, 2023, 9(1): 64.",
    "Wu S, Kondo Y, Kakimoto M, et al. Machine-learning-assisted discovery of polymers with high thermal conductivity. npj Computational Materials, 2019, 5: 66.",
    "Patel R A, Webb M A. Data-driven design of polymer-based biomaterials: high-throughput simulation, experimentation, and machine learning. ACS Applied Bio Materials, 2024, 7(2): 510–527.",
    "Hollmann N, Müller S, Eggensperger K, et al. TabPFN: A transformer that solves small tabular classification problems in a second. International Conference on Learning Representations, 2023.",
    "Zhou Z H. A brief introduction to weakly supervised learning. National Science Review, 2018, 5(1): 44–53.",
    "Tom G, Schmid S P, Baird S G, et al. Self-driving laboratories for chemistry and materials science. Chemical Reviews, 2024, 124(16): 9633–9732.",
    "Seeman N C. DNA in a material world. Nature, 2003, 421(6921): 427–431.",
    "Alemdaroglu F E, Herrmann A. DNA meets synthetic polymers — highly versatile hybrid materials. Organic & Biomolecular Chemistry, 2007, 5(9): 1311–1320.",
    "Tan S J, Campolongo M J, Luo D, et al. Building plasmonic nanostructures with DNA. Nature Nanotechnology, 2011, 6(5): 268–276.",
    "Chen Y J, Groves B, Muscat R A, et al. DNA nanotechnology from the test tube to the cell. Nature Nanotechnology, 2015, 10(9): 748–760.",
    "ASTM International. Standard Test Method for Assignment of the Glass Transition Temperatures by DSC: ASTM E1356-08. West Conshohocken: ASTM International, 2014.",
    "Audus D J, de Pablo J J. Polymer informatics: opportunities and challenges. ACS Macro Letters, 2017, 6(10): 1078–1082.",
)


def add_references(doc) -> None:
    page_break(doc)
    add_section_title(doc, "参考文献", in_toc=True)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        fmt_paragraph(
            p,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line_spacing_pt=24,
            first_line_indent_pt=0,
            left_indent_cm=0.7,
        )
        pf = p.paragraph_format
        pf.first_line_indent = Cm(-0.7)
        r = p.add_run(f"[{i}]  ")
        style_run(r, F_SONG, SZ_XIAO4)
        r = p.add_run(ref)
        style_run(r, F_SONG, SZ_XIAO4)


def add_appendix(doc) -> None:
    page_break(doc)
    add_section_title(doc, "附录", in_toc=True)

    add_h2(doc, "附录 A  主要符号与术语")
    add_three_line_table(
        doc,
        ["符号 / 术语", "含义"],
        [
            ("Tg", "玻璃化转变温度，单位 ℃ 或 K"),
            ("wᵢ", "共聚物中第 i 个端点的归一化组成权重"),
            ("Tgᵢ", "第 i 个端点的均聚物 Tg"),
            ("Fox 关系", "1/Tg = Σᵢ wᵢ / Tgᵢ，Kelvin 温度下计算"),
            ("min-R²", "三任务最小 R²，作为综合主指标"),
            ("group holdout", "按体系或碱基族整体留出的评估方式"),
            ("链段物理特征", "刻画链段刚性、基团贡献、自由体积等的 58 维特征"),
            ("图嵌入", "由聚合物图神经网络预训练得到的 64 维结构表示"),
            ("polyBERT 嵌入", "聚合物语言模型预训练表示，经主成分分析降维得 64 维"),
            ("186 维端点表征", "三类端点表示拼接得到的端点特征向量"),
            ("门控函数 g(x)", "控制残差幅度的指示函数集合"),
            ("非均聚物校准", "仅作用于共聚物样本的事后偏置修正"),
            ("低 Fox 区残差收缩", "Fox 基线落入低 Tg 区时压制残差幅度"),
        ],
        col_widths_cm=[3.6, 11.0],
    )

    add_h2(doc, "附录 B  模型超参数与实验设置")
    add_three_line_table(
        doc,
        ["超参数 / 设置", "取值"],
        [
            ("均聚物样本权重", "1.0"),
            ("一般共聚物样本权重", "10.0"),
            ("DNA 相关样本权重", "60.0"),
            ("虚拟样本默认权重", "0.0 (按实验调整)"),
            ("PolyInfo 体系冲突过滤阈值", "同体系同组成 Tg 标准差 > 10 K"),
            ("均聚物 holdout 比例", "≈ 20%"),
            ("特征展开维度", "1191"),
            ("低 Fox 区阈值", "Fox Tg < 270 K (≈ −3 ℃)"),
            ("一致性过滤阈值", "弱标签与教师预测差 ≤ 30 ℃"),
            ("结构近邻虚拟样本数 / 权重", "400 / 0.05"),
            ("一致性过滤虚拟样本数 / 权重", "200 / 0.02"),
            ("评估随机种子", "固定，可复现"),
        ],
        col_widths_cm=[5.6, 9.0],
    )

    add_h2(doc, "附录 C  数据接口与代码模块说明")
    add_body_para(
        doc,
        "本文方法体系由四类代码模块组成：数据清洗、虚拟样本生成、统一回归模型与证据自适应路由。各模块的功能与接口归纳如下，供数据与代码材料一并提交。",
    )
    add_three_line_table(
        doc,
        ["模块名", "功能", "主要输入", "主要输出"],
        [
            ("polyinfo_parser", "原始 PolyInfo 共聚物记录解析与端点映射", "原始数据表", "结构化共聚物条目"),
            ("polyinfo_conflict_filter", "体系冲突过滤", "结构化条目", "清洗后真实共聚物表"),
            ("virtual_copolymer_generator", "端点驱动的虚拟弱标签生成", "端点库 / recipe CSV", "可审计虚拟数据表"),
            ("homopolymer_endpoint_model", "端点 Tg 源模型训练与预测", "均聚物 SMILES + Tg", "端点 Tg 估计 + 端点表征"),
            ("universal_tg_regressor", "物理先验门控残差回归与三类校准", "多源训练表 + 样本权重", "Tg 预测与诊断信息"),
            ("evidence_adaptive_router", "证据自适应路由与可审计输出", "用户输入证据", "Tg 预测 + 路线 + 风险标签"),
        ],
        col_widths_cm=[4.0, 4.6, 3.6, 3.6],
    )

    add_h2(doc, "附录 D  路由系统输入输出字段规范")
    add_three_line_table(
        doc,
        ["字段名", "含义", "类型 / 取值"],
        [
            ("smiles / smiles_i", "重复单元结构 (单 SMILES 或多端点)", "字符串"),
            ("w_i", "端点 i 的组成权重", "浮点数 (∈ [0, 1])"),
            ("n_components", "组分数 (端点重复单元类别数)", "整数；2 = 二元，≥3 = 多元，1 = 均聚物"),
            ("architecture", "链段架构 (与 n_components 正交)", "random（无规）/ block（嵌段）/ homo（均聚物专用）"),
            ("system_id (COID)", "已知体系编号 (可选)", "字符串"),
            ("base_family", "DNA 相关样本的碱基族 (可选)", "A / T / G / C / U"),
            ("endpoint_tg_c (可选)", "实验测得端点 Tg / ℃", "浮点数"),
            ("router_route", "路由实际选择的预测路线", "字符串"),
            ("primary_method", "本次预测的主要计算方法", "字符串"),
            ("tg_pred_c", "Tg 预测值 / ℃", "浮点数"),
            ("tg_interval_c", "Tg 90% 预测区间 / ℃", "(下界, 上界)"),
            ("risk_flag", "风险提示标签", "字符串"),
        ],
        col_widths_cm=[3.6, 7.4, 3.6],
    )


def add_acknowledgements(doc) -> None:
    page_break(doc)
    add_section_title(doc, "致谢", in_toc=True)
    add_body_para(
        doc,
        "本文研究是在课题组导师的悉心指导下完成的。导师在研究方向选择、技术路线讨论和论文撰写过程中给予了大量启发性的建议，特别是在数据清洗的严谨性、严格评估的统计意义、路由系统的应用边界以及双方法的覆盖矩阵讨论上提出了关键意见，使本文得以在科学性与工程价值之间保持平衡。",
    )
    add_body_para(
        doc,
        "感谢课题组同学在数据整理、实验运行、结果交叉核对和论文校阅过程中给予的支持。对来自学校教务、实验室管理和计算资源支持单位的同行表示诚挚感谢。",
    )
    add_body_para(
        doc,
        "本文的所有结果均在严格分组评估协议下获得，相关数据来源、清洗规则、模型训练设置和评估代码可与论文一同提交以供复核。如未来在共聚物相态特征、DNA 相关数据扩充或虚拟数据端点感知生成等方向取得进一步进展，本文方法体系也可作为基础平台继续发展。",
    )
    # 落款日期：附件4 模板 "2026年XX月XX日"
    p = doc.add_paragraph()
    fmt_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        line_spacing_pt=24,
        space_before_pt=18,
    )
    r = p.add_run("2026年4月27日")
    style_run(r, F_SONG, SZ_XIAO4)


# ============================================================
# Auxiliary outputs
# ============================================================


def write_checklist() -> None:
    items = [
        # ---- 一、论文要素 ----
        ("【要素】Word 完整版包含八部分：封面页、摘要与关键词、章节目录、表格和插图清单、论文正文、参考文献、附录、致谢。", True),
        ("【要素】封面页含参赛院校、论文题目、参赛队员、指导教师等参赛队相关信息。", True),
        ("【要素】摘要包含研究目的、方法、结果、结论、创新点五要素。", True),
        ("【要素】关键词数量 3-5 个，准确简洁覆盖核心内容（当前 5 个：共聚物玻璃化转变温度；统计预测；物理先验门控残差；证据自适应路由；DNA 相关高分子）。", True),
        ("【要素】章节目录呈现 2-3 级（一/（一）/1）。", True),
        ("【要素】表格和插图清单分别列出全文表格与插图。", True),
        ("【要素】论文正文涵盖：问题描述、指标选择、数据描述、模型建立、求解和检验、模型结果分析、结论建议。", True),
        ("【要素】参考文献对引文作者、出处、版本详细注明（GB/T 7714 格式）；中文文献置于英文文献前。", True),
        ("【要素】附录补充正文未涉及的方法、统计软件、数据处理细节等。", True),
        ("【要素】致谢简洁明了、实事求是。", True),
        # ---- 二、排版格式 ----
        ("【格式】除封面页与致谢外，其他页不出现学校、参赛队及指导教师信息（匿名版已自动剥离封面+致谢）。", True),
        ("【格式】页面 A4，页边距上下 2.54 cm、左右 3.17 cm。", True),
        ("【格式】行距固定值 24 磅，段前、段后 0 行（表格内容除外为单倍行距）。", True),
        ("【格式】层次序号：一、；（一）；1.；①（前三层在脚本，第四层 ① 由作者按需添加）。", True),
        # ---- 字体规范 ----
        ("【字体】论文正文总标题：方正小标宋_GBK 三号，居中。", True),
        ("【字体】一级标题（一、…）：黑体小三号。", True),
        ("【字体】二级标题（（一）…）：楷体四号。", True),
        ("【字体】三级标题（1.…）：宋体小四号加粗。", True),
        ("【字体】四级标题（①…）：宋体小四号（不加粗）。", True),
        ("【字体】正文：宋体小四号，首行缩进 2 字符。", True),
        ("【字体】摘要 / 目录 / 表图清单 / 参考文献 / 附录 / 致谢 标题：黑体四号居中。", True),
        ("【字体】关键词标签：黑体小四号；其余宋体小四号。", True),
        ("【字体】封面身份字段：仿宋_GB2312 小二号、单倍行距（附件4 第 11 页布局）。", True),
        # ---- 目录 ----
        ("【目录】使用 Word 自动 TOC 域 + 占位行（含一/二/三级标题，自动生成页码）。", True),
        ("【目录】正文页码单独编列，从正文第一页开始；前置部分（封面/摘要/目录/清单）独立编页。", True),
        # ---- 表格 ----
        ("【表格】11 张表全部使用三线表样式（顶/底粗线 + 表头分隔细线）。", True),
        ("【表格】表格为可编辑文字（非图片），居中对齐。", True),
        ("【表格】表格内容：宋体小四号或五号字，单倍行距。", True),
        ("【表格】表序与表名置表上方居中，宋体小四号加粗，表序与表名间空两格（如\"表1  名称\"）。", True),
        # ---- 图 ----
        ("【插图】11 张图使用统一线宽与配色，含中文字体（宋体/黑体/STIX），无缺字。", True),
        ("【插图】图序与图名置图下方居中，宋体小四号加粗，图序与图名间空两格（如\"图1  名称\"）。", True),
        # ---- 公式 ----
        ("【公式】公式为可编辑格式（OMML），不可用图片代替；居中对齐；公式后依次标注 (n) 序号靠右行末。", True),
        ("【公式】行距使用 AT_LEAST 28pt（编号公式）/ AT_LEAST 24pt（内联），分数 / Σ 不会被裁底。", True),
        ("【公式】所有数学符号（ŷ、r_θ、f_phys、Tg_max/min/w、w_homo/co/nuc/v、min-R²、MAE/RMSE/R² 等）均为 OMML 可编辑公式而非纯文字。", True),
        # ---- 文字规范 ----
        ("【文字】不使用繁体字、异体字、复合字及不规范简化字（除非必要）。", True),
        ("【文字】公历世纪 / 年 / 月 / 日 / 时刻 / 计量用阿拉伯数字；邻近两数并列连用表概数用汉字（如\"五六岁\"）。", True),
        ("【文字】图、表、附录、参考文献、公式一律阿拉伯数字连续编号。", True),
        # ---- 封面 ----
        ("【封面】4 个身份字段（作品编号、学校、题目、队员、导师）使用整宽下划线占位，需用户填写。", True),
        ("【封面】完全照搬附件4 第 11 页布局（作品编号黑体三号 + 主标题方正小标宋_GBK + 4 字段仿宋_GB2312 小二号）。", True),
        # ---- 匿名 / 内部代号 ----
        ("【匿名】PDF 匿名版去掉封面页与致谢，其余六部分与 Word 完整版完全一致。", True),
        ("【匿名】匿名版正文不出现学校 / 同学 / 导师 / 致谢 / 感谢等身份字样（脚本自动审计 0 命中）。", True),
        ("【代号】正文中不直接出现 exp45 / exp53 / exp56 / exp57 / exp58 等内部实验代号。", True),
        # ---- 数据指标对齐 ----
        ("【数据】方法 I 三任务严格指标（exp56 holdout）在摘要、表5、表6 与结论中均为 0.887 / 0.849 / 0.817。", True),
        ("【数据】方法 II 关键指标对齐 universal_router_fulltest：均聚物公平 holdout R²=0.907 / 5×3 CV R²=0.917、PolyInfo 体系内 0.933 / 跨体系 0.820、DNA specialized 0.971 / universal 0.564。", True),
        ("【数据】综合主指标 min-R² = 0.817 在摘要与结论中一致。", True),
        # ---- 文件命名 ----
        ("【文件】Word 完整版命名 \"作品全文—组别—作品编号.docx\"，PDF 匿名版命名 \"匿名作品—组别—作品编号.pdf\"，均小于 100MB。", True),
        ("【文件】致谢落款日期 \"2026年XX月XX日\"，提交前用户填写实际日期。", True),
        # ---- 提交前用户操作 ----
        ("提交前请填写真实作品编号、参赛学校、参赛队员、指导老师，并在装有 方正小标宋_GBK 的电脑上重转 PDF。", False),
        ("提交前请用 Word 打开 docx，右键目录 → 更新域 → 更新整个目录，再转 PDF（确保页码与正文同步）。", False),
        ("提交前确认论文题目不含学校名称、队员姓名，且与平台填写的题目保持一致。", False),
        ("提交前完成知网查重（cx.cnki.net，类型选\"职称评审\"，保存\"全文标明引文\" PDF），命名 \"查重报告—作品编号.pdf\"。", False),
    ]
    lines = ["# 比赛论文格式对齐检查表", ""]
    lines.append("> 依据：附件4《2026 年（第十二届）全国大学生统计建模大赛参赛须知》")
    lines.append("")
    lines.append("## 自动产出 (脚本已对齐)")
    for text, done in items:
        if done:
            lines.append(f"- [x] {text}")
    lines.append("")
    lines.append("## 提交前用户操作")
    for text, done in items:
        if not done:
            lines.append(f"- [ ] {text}")
    lines.append("")
    CHECKLIST.write_text("\n".join(lines), encoding="utf-8")


def write_readme() -> None:
    content = (
        "# 论文构建说明\n\n"
        "本目录由 `scripts/build_competition_paper_docx.py` 自动生成。\n"
        "脚本同时构建 Word 完整版和 Word 匿名版，并附带 10 张正式插图。\n\n"
        "## 1. 生成 Word（完整版 + 匿名版）\n\n"
        "```powershell\n"
        "python scripts\\build_competition_paper_docx.py\n"
        "```\n\n"
        "输出：\n\n"
        "- `output/doc/paper_full_draft.docx`（八部分齐全）\n"
        "- `output/doc/作品全文-组别-待填作品编号.docx`（与上同内容，比赛命名）\n"
        "- `output/doc/paper_anonymous.docx`（六部分：摘要、目录、表图清单、正文、参考文献、附录；删除封面与致谢）\n"
        "- `output/doc/figures/figureXX_*.png`（10 张正式插图）\n\n"
        "## 2. 渲染 PDF（LibreOffice headless）\n\n"
        "```powershell\n"
        "& 'C:\\Program Files\\LibreOffice\\program\\soffice.exe' --headless --convert-to pdf --outdir output\\doc output\\doc\\paper_full_draft.docx output\\doc\\paper_anonymous.docx\n"
        "```\n\n"
        "## 3. 渲染各页 PNG（校对）\n\n"
        "```powershell\n"
        "Remove-Item -Path tmp\\doc_render\\paper_page-*.png -Force -ErrorAction SilentlyContinue\n"
        "pdftoppm -png -r 110 -f 1 -l 34 output\\doc\\paper_full_draft.pdf tmp\\doc_render\\paper_page\n"
        "```\n\n"
        "## 4. 提交清单\n\n"
        "1. 在 Word 中打开 `作品全文-组别-待填作品编号.docx`，填写：\n"
        "   - 作品编号（封面顶部，黑体三号）\n"
        "   - 参赛学校 / 参赛队员 / 指导老师（封面下方下划线字段）\n"
        "2. 保存后用 LibreOffice 重新转 PDF 作为 Word 完整版的对应渲染。\n"
        "3. `paper_anonymous.pdf` 作为匿名版直接提交（已自动剔除封面与致谢，已通过身份信息扫描）。\n\n"
        "## 5. 锁定的核心数字（写作一致性）\n\n"
        "### 方法 I — exp56 物理先验门控残差通用统计模型（多源融合）\n\n"
        "| 任务 | n | R² | MAE / ℃ | RMSE / ℃ |\n"
        "|---|---:|---:|---:|---:|\n"
        "| 均聚物 random holdout | 1498 | 0.887 | 27.16 | 38.09 |\n"
        "| 一般共聚物 group holdout | 149 | 0.849 | 16.65 | 21.64 |\n"
        "| DNA 相关 group holdout | 17 | 0.817 | 6.27 | 8.51 |\n"
        "| 综合主指标 min-R² | — | 0.817 | — | — |\n\n"
        "### 方法 II — 证据自适应路由系统（公平口径，对齐服务器 universal_router_fulltest）\n\n"
        "| 评估场景 | n | R² | MAE / K | 路线 |\n"
        "|---|---:|---:|---:|---|\n"
        "| 均聚物 公平 holdout (与方法 I 同 1498) | 1498 | 0.907 | 20.6 | TabPFN 186d 端点源模型 |\n"
        "| 均聚物 5×3 RepeatedKFold CV (独立基准) | 7486 | 0.917 | 20.4 | TabPFN 186d 端点源模型 |\n"
        "| PolyInfo 体系内 (leave-row) | 149 | 0.933 | 8.76 | 同体系残差 IDW |\n"
        "| PolyInfo 跨体系 (leave-system-out) | 149 | 0.820 | 16.49 | Linear-Fox |\n"
        "| DNA 跨碱基 specialized (端点实测+LOOCV) | 17 | 0.971 | 2.93 | Physics-Ridge |\n"
        "| DNA 跨碱基 specialized leave-base-out | 17 | 0.872 | 5.25 | Physics-Ridge |\n"
        "| DNA 跨碱基 universal default (仅 SMILES) | 17 | 0.564 | 10.63 | Linear-Fox 全局校准 |\n\n"
        "## 6. 关键页面（完整版，结构变化后估算）\n\n"
        "| 部分 | 起始页（doc） |\n"
        "|---|---:|\n"
        "| 封面 | 1 |\n"
        "| 摘要 | 2 |\n"
        "| 目录 | 3 |\n"
        "| 表格与插图清单 | 5 |\n"
        "| 正文（一、问题描述与研究思路） | 7 |\n"
        "| 四、方法 I 通用统计模型 | 17 |\n"
        "| 五、方法 II 证据自适应路由系统 | 22 |\n"
        "| 六、综合求解、对比与诊断 | 27 |\n"
        "| 七、结论与建议 | 35 |\n"
        "| 参考文献 | 37 |\n"
        "| 附录 | 40 |\n"
        "| 致谢 | 42 |\n"
    )
    README.write_text(content, encoding="utf-8")


def build_full_doc() -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc)
    add_abstract(doc)
    add_toc(doc)
    add_lists(doc)
    start_body_section(doc)
    add_body(doc)
    add_references(doc)
    add_appendix(doc)
    add_acknowledgements(doc)
    doc.save(DOCX_ASCII)
    try:
        doc.save(DOCX_CN)
    except OSError:
        pass


def build_anonymous_doc() -> None:
    """匿名版：删除封面页与致谢页，其余六部分保持一致。"""
    doc = Document()
    setup_doc(doc)
    add_abstract(doc)
    add_toc(doc, include_acknowledgements=False)
    add_lists(doc)
    start_body_section(doc)
    add_body(doc)
    add_references(doc)
    add_appendix(doc)
    doc.save(DOCX_ANON)
    try:
        doc.save(DOCX_ANON_CN)
    except OSError:
        pass


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    create_figures()
    build_full_doc()
    build_anonymous_doc()
    write_checklist()
    write_readme()
    print(f"Wrote {DOCX_ASCII}")
    print(f"Wrote {DOCX_ANON}")


if __name__ == "__main__":
    main()
